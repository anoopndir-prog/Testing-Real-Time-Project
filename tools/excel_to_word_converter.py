#!/usr/bin/env python3
"""Convert a filled SKF Excel test request into a project specification Word document."""

from __future__ import annotations

import argparse
import datetime as dt
import io
import math
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Dict, Iterable, Optional, Tuple

import openpyxl
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from openpyxl.styles.colors import COLOR_INDEX
from openpyxl.utils import get_column_letter, range_boundaries
from PIL import Image, ImageDraw, ImageFont


DEFAULT_FONT_PATHS = [
    "/System/Library/Fonts/Supplemental/Arial.ttf",
    "/System/Library/Fonts/Supplemental/Calibri.ttf",
]


@dataclass
class ExtractedData:
    requester: str
    customer: str
    application: str
    purpose: str
    num_samples: str
    part_number: str
    test_type: str

    shaft_diameter: str
    shaft_tolerance: str
    shaft_unit: str
    shaft_material: str
    shaft_ra: str
    shaft_hardness: str

    bore_diameter: str
    bore_tolerance: str
    bore_unit: str
    bore_material: str
    bore_ra: str

    dro_value: str
    dro_tolerance: str
    stbm_value: str
    stbm_tolerance: str
    seal_cock_value: str
    seal_cock_tolerance: str
    recip_value: str
    recip_tolerance: str

    oil_type: str
    pre_lube_required: str
    setup_notes: str
    setup_notes_full: str

    contamination_type: str
    contamination_mix_ratio: str
    contamination_amount: str
    contamination_recip_freq: str
    contamination_stbm_orientation: str

    oil_change_interval: str
    acceptance_duration: str
    acceptance_failure: str


class TemplateUpdateError(RuntimeError):
    pass


def _load_font(size: int) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    for font_path in DEFAULT_FONT_PATHS:
        candidate = Path(font_path)
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size)
    return ImageFont.load_default()


def _fmt_num(value: Any, decimals: int = 3) -> str:
    if value is None or value == "":
        return ""
    if isinstance(value, str):
        text = value.strip()
        return text
    if isinstance(value, (int, float)):
        return f"{float(value):.{decimals}f}"
    return str(value)


def _fmt_general(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, dt.datetime):
        return value.strftime("%Y-%m-%d")
    if isinstance(value, float):
        text = f"{value:.6f}".rstrip("0").rstrip(".")
        return text if text else "0"
    return str(value).strip()


def _cell(ws: openpyxl.worksheet.worksheet.Worksheet, ref: str, default: str = "") -> str:
    value = ws[ref].value
    if value is None:
        return default
    return _fmt_general(value)


def _normalize_mix_ratio(value: str) -> str:
    text = (value or "").strip()
    if not text:
        return text
    parts = text.split(":")
    if len(parts) == 3:
        h, m, s = parts
        if s in {"00", "0"}:
            try:
                return f"{int(h)}:{int(m)}"
            except ValueError:
                return text
    return text


def _normalize_ra(value: str) -> str:
    text = (value or "").strip()
    if not text:
        return ""
    lowered = text.lower()
    if lowered.startswith("ra"):
        text = text[2:].strip()
    text = text.replace("-", "~")
    text = " ".join(text.split())
    return f"{text} µm Ra" if text else ""


def _is_contamination_test(test_type: str, purpose: str) -> bool:
    t = (test_type or "").lower()
    p = (purpose or "").lower()
    triggers = ["mud", "slurry", "dry dust", "contamination"]
    return any(token in t for token in triggers) or any(token in p for token in triggers)


def _collect_setup_notes(ws: openpyxl.worksheet.worksheet.Worksheet) -> str:
    notes: list[str] = []

    for row in range(60, 66):
        for col in range(2, 13):  # B:L
            val = ws.cell(row, col).value
            txt = _safe_text(val)
            if not txt:
                continue
            lowered = txt.lower()
            if lowered in {"setup notes", "notes"}:
                continue
            notes.append(txt)

    # preserve order and remove duplicates
    deduped: list[str] = []
    seen = set()
    for item in notes:
        if item not in seen:
            deduped.append(item)
            seen.add(item)
    return "\n".join(deduped)


def extract_excel_data(excel_path: Path) -> ExtractedData:
    wb = openpyxl.load_workbook(excel_path, data_only=True)
    ws1 = wb["Page 1"]
    ws2 = wb["Page 2"]

    test_type = _cell(ws1, "K10")
    purpose = _cell(ws1, "C11")
    setup_notes_full = _collect_setup_notes(ws1)

    return ExtractedData(
        requester=_cell(ws1, "C4"),
        customer=_cell(ws1, "G4"),
        application=_cell(ws1, "C10"),
        purpose=purpose,
        num_samples=_cell(ws1, "L6"),
        part_number=_cell(ws1, "C6"),
        test_type=test_type,
        shaft_diameter=_fmt_num(ws1["C18"].value, 3),
        shaft_tolerance=_fmt_num(ws1["F18"].value, 3),
        shaft_unit=_cell(ws1, "D18", "mm"),
        shaft_material=_cell(ws1, "H18"),
        shaft_ra=_cell(ws1, "K18"),
        shaft_hardness=_cell(ws1, "L18"),
        bore_diameter=_fmt_num(ws1["C20"].value, 3),
        bore_tolerance=_fmt_num(ws1["F20"].value, 3),
        bore_unit=_cell(ws1, "D20", "mm"),
        bore_material=_cell(ws1, "H20"),
        bore_ra=_cell(ws1, "K20"),
        dro_value=_fmt_num(ws1["D50"].value, 3),
        dro_tolerance=_fmt_num(ws1["G50"].value, 3),
        stbm_value=_fmt_num(ws1["D52"].value, 3),
        stbm_tolerance=_fmt_num(ws1["G52"].value, 3),
        seal_cock_value=_fmt_num(ws1["D54"].value, 3),
        seal_cock_tolerance=_fmt_num(ws1["G54"].value, 3),
        recip_value=_fmt_num(ws1["D56"].value, 3),
        recip_tolerance=_fmt_num(ws1["G56"].value, 3),
        oil_type=_cell(ws1, "B48"),
        pre_lube_required=_cell(ws1, "J48"),
        setup_notes=_cell(ws1, "B61"),
        setup_notes_full=setup_notes_full,
        contamination_type=_cell(ws1, "J50"),
        contamination_mix_ratio=_normalize_mix_ratio(_cell(ws1, "J52")),
        contamination_amount=_cell(ws1, "J54"),
        contamination_recip_freq=_cell(ws1, "J56"),
        contamination_stbm_orientation=_cell(ws1, "J58"),
        oil_change_interval=_cell(ws2, "E4"),
        acceptance_duration=_cell(ws2, "G5"),
        acceptance_failure=_cell(ws2, "G6"),
    )


def _safe_text(value: Any) -> str:
    text = _fmt_general(value)
    return text.replace("\n", " ").strip()


def _iter_nonempty_rows(
    ws: openpyxl.worksheet.worksheet.Worksheet, min_row: int, max_row: int, min_col: int, max_col: int
) -> Iterable[int]:
    for row in range(min_row, max_row + 1):
        has_any = False
        for col in range(min_col, max_col + 1):
            if _safe_text(ws.cell(row, col).value):
                has_any = True
                break
        if has_any:
            yield row


def _argb_to_rgb(argb: str) -> Optional[Tuple[int, int, int]]:
    if not argb:
        return None
    val = argb.strip().lstrip("#")
    if len(val) == 8:
        val = val[2:]
    if len(val) != 6:
        return None
    try:
        return (int(val[0:2], 16), int(val[2:4], 16), int(val[4:6], 16))
    except ValueError:
        return None


def _resolve_openpyxl_color(color_obj, default_rgb: Tuple[int, int, int]) -> Tuple[int, int, int]:
    if color_obj is None:
        return default_rgb

    color_type = getattr(color_obj, "type", None)
    if color_type == "rgb":
        rgb = _argb_to_rgb(getattr(color_obj, "rgb", ""))
        return rgb or default_rgb
    if color_type == "indexed":
        idx = getattr(color_obj, "indexed", None)
        if isinstance(idx, int) and 0 <= idx < len(COLOR_INDEX):
            rgb = _argb_to_rgb(COLOR_INDEX[idx])
            return rgb or default_rgb

    # indexed/theme colors are not directly resolved by openpyxl without theme mapping;
    # keep default when unresolved.
    return default_rgb


def _load_cell_font(font_name: str, size: int, bold: bool = False, italic: bool = False):
    size = max(8, int(size))
    candidates = []
    if font_name:
        candidates.extend(
            [
                f"/System/Library/Fonts/Supplemental/{font_name}.ttf",
                f"/System/Library/Fonts/Supplemental/{font_name}.TTF",
            ]
        )
    candidates.extend(DEFAULT_FONT_PATHS)
    for path in candidates:
        p = Path(path)
        if p.exists():
            try:
                return ImageFont.truetype(str(p), size=size)
            except Exception:
                continue
    return ImageFont.load_default()


def render_sheet_range_to_image(
    ws: openpyxl.worksheet.worksheet.Worksheet,
    range_ref: str,
    out_path: Path,
    end_at_last_content: bool = False,
) -> None:
    zoom = 1.65
    min_col, min_row, max_col, max_row = range_boundaries(range_ref)

    if end_at_last_content:
        rows = list(_iter_nonempty_rows(ws, min_row, max_row, min_col, max_col))
        if rows:
            max_row = max(rows)

    default_col_width = 8.43
    default_row_height = 15.0

    col_widths = []
    for col in range(min_col, max_col + 1):
        width = ws.column_dimensions[get_column_letter(col)].width
        width = width if width else default_col_width
        col_widths.append(max(24, int((width * 7 + 6) * zoom)))

    row_heights = []
    for row in range(min_row, max_row + 1):
        height = ws.row_dimensions[row].height
        height = height if height else default_row_height
        row_heights.append(max(18, int((height * 96 / 72) * zoom)))

    margin_x = 0
    margin_y = 0

    total_width = sum(col_widths) + margin_x * 2
    total_height = sum(row_heights) + margin_y * 2

    img = Image.new("RGB", (total_width, total_height), color="white")
    draw = ImageDraw.Draw(img)
    default_font = _load_font(max(11, int(11 * zoom)))

    merged_top_left: Dict[Tuple[int, int], Tuple[int, int]] = {}
    merged_covered: Dict[Tuple[int, int], Tuple[int, int]] = {}
    for merged in ws.merged_cells.ranges:
        m_min_col, m_min_row, m_max_col, m_max_row = range_boundaries(str(merged))
        if m_max_col < min_col or m_min_col > max_col or m_max_row < min_row or m_min_row > max_row:
            continue
        top_left = (m_min_row, m_min_col)
        merged_top_left[top_left] = (m_max_row, m_max_col)
        for rr in range(m_min_row, m_max_row + 1):
            for cc in range(m_min_col, m_max_col + 1):
                if (rr, cc) != top_left:
                    merged_covered[(rr, cc)] = top_left

    x_offsets = [margin_x]
    for w in col_widths:
        x_offsets.append(x_offsets[-1] + w)

    y_offsets = [margin_y]
    for h in row_heights:
        y_offsets.append(y_offsets[-1] + h)

    default_grid = (160, 160, 160)
    font_cache: Dict[Tuple[str, int, bool, bool], Any] = {}

    for r_idx, row in enumerate(range(min_row, max_row + 1)):
        for c_idx, col in enumerate(range(min_col, max_col + 1)):
            if (row, col) in merged_covered:
                continue

            cell = ws.cell(row, col)
            x0 = x_offsets[c_idx]
            y0 = y_offsets[r_idx]
            x1 = x_offsets[c_idx + 1]
            y1 = y_offsets[r_idx + 1]

            if (row, col) in merged_top_left:
                m_max_row, m_max_col = merged_top_left[(row, col)]
                m_max_row = min(m_max_row, max_row)
                m_max_col = min(m_max_col, max_col)
                m_r_idx = m_max_row - min_row + 1
                m_c_idx = m_max_col - min_col + 1
                x1 = x_offsets[m_c_idx]
                y1 = y_offsets[m_r_idx]

            fill_rgb = (255, 255, 255)
            pattern = (cell.fill.patternType or "").lower() if cell.fill else ""
            if pattern and pattern != "none":
                fill_rgb = _resolve_openpyxl_color(getattr(cell.fill, "fgColor", None), (255, 255, 255))
            draw.rectangle([x0, y0, x1, y1], fill=fill_rgb)

            border_color = default_grid
            draw.rectangle([x0, y0, x1, y1], outline=border_color, width=1)

            value = _safe_text(cell.value)
            if value:
                max_text_width = max(20, x1 - x0 - 8)
                words = value.replace("\n", " ").split(" ")
                lines = []
                current = ""
                for word in words:
                    test = f"{current} {word}".strip()
                    cell_font = cell.font
                    font_name = (cell_font.name or "").strip()
                    font_size = int((int(cell_font.sz) if cell_font and cell_font.sz else 11) * zoom)
                    is_bold = bool(cell_font.bold) if cell_font else False
                    is_italic = bool(cell_font.italic) if cell_font else False
                    key = (font_name, font_size, is_bold, is_italic)
                    if key not in font_cache:
                        font_cache[key] = _load_cell_font(font_name, font_size, is_bold, is_italic)
                    font = font_cache.get(key, default_font)

                    bbox = draw.textbbox((0, 0), test, font=font)
                    if bbox[2] - bbox[0] <= max_text_width:
                        current = test
                    else:
                        if current:
                            lines.append(current)
                        current = word
                if current:
                    lines.append(current)

                text_color = _resolve_openpyxl_color(getattr(cell.font, "color", None), (20, 20, 20))
                align = (cell.alignment.horizontal or "").lower() if cell.alignment else ""
                line_height = max(14, int((font.size if hasattr(font, "size") and font.size else 12) * 1.25))

                for line_idx, line in enumerate(lines[:8]):
                    ty = y0 + 3 + line_idx * line_height
                    if ty + line_height > y1:
                        break

                    text_w = draw.textbbox((0, 0), line, font=font)[2]
                    if align in {"center", "centercontinuous"}:
                        tx = x0 + max(2, int((x1 - x0 - text_w) / 2))
                    elif align in {"right", "distributed", "justify"}:
                        tx = x1 - text_w - 4
                    else:
                        tx = x0 + 4
                    draw.text((tx, ty), line, fill=text_color, font=font)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    img.save(out_path)


def _set_paragraph_text(paragraph, new_text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(new_text)
        return
    paragraph.runs[0].text = new_text
    for run in paragraph.runs[1:]:
        run.text = ""


def _set_cell_text(cell, text: str) -> None:
    if not cell.paragraphs:
        p = cell.add_paragraph("")
    else:
        p = cell.paragraphs[0]
    _set_paragraph_text(p, text)
    for extra in cell.paragraphs[1:]:
        _set_paragraph_text(extra, "")


def _copy_font_style(from_run, to_run) -> None:
    to_run.font.name = from_run.font.name
    to_run.font.size = from_run.font.size
    to_run.bold = from_run.bold
    to_run.italic = from_run.italic
    to_run.underline = from_run.underline
    to_run.font.color.rgb = from_run.font.color.rgb


def _set_paragraph_font_like(paragraph, source_paragraph) -> None:
    if not paragraph.runs or not source_paragraph or not source_paragraph.runs:
        return
    src_run = None
    for run in source_paragraph.runs:
        if run.text:
            src_run = run
            break
    if src_run is None:
        src_run = source_paragraph.runs[0]
    for run in paragraph.runs:
        _copy_font_style(src_run, run)


def _find_paragraph_by_prefix(doc: Document, prefix: str, occurrence: int = 1):
    count = 0
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            count += 1
            if count == occurrence:
                return p
    raise TemplateUpdateError(f"Could not find paragraph starting with: {prefix!r}")


def _replace_image_after_anchor(doc: Document, anchor_text: str, image_path: Path, width_inches: float = 6.5) -> None:
    anchor_index: Optional[int] = None
    for idx, para in enumerate(doc.paragraphs):
        if anchor_text in para.text:
            anchor_index = idx
            break

    if anchor_index is None:
        raise TemplateUpdateError(f"Could not locate anchor text: {anchor_text!r}")

    target = None
    for idx in range(anchor_index + 1, len(doc.paragraphs)):
        para = doc.paragraphs[idx]
        has_drawing = any("drawing" in run._element.xml for run in para.runs)
        if has_drawing:
            target = para
            break

    if target is None:
        target = doc.paragraphs[anchor_index + 1]

    for run in list(target.runs):
        target._p.remove(run._element)

    run = target.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))
    target.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _extract_main_drawing_image(
    wb: openpyxl.Workbook,
    out_path: Path,
    sheet_name: str = "Part Drawing 1",
) -> Optional[Path]:
    if sheet_name not in wb.sheetnames:
        return None

    ws = wb[sheet_name]
    images = list(getattr(ws, "_images", []) or [])
    if not images:
        return None

    # Pick the largest embedded image (typically the main drawing block).
    chosen = None
    chosen_area = -1
    for img in images:
        try:
            blob = img._data()
            with Image.open(io.BytesIO(blob)) as im:
                area = im.width * im.height
            if area > chosen_area:
                chosen_area = area
                chosen = blob
        except Exception:
            continue

    if not chosen:
        return None

    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_bytes(chosen)
    return out_path


def _replace_acceptance_line(doc: Document, acceptance_text: str) -> None:
    anchor_index: Optional[int] = None
    for idx, para in enumerate(doc.paragraphs):
        if para.text.strip().startswith("Acceptance Criteria"):
            anchor_index = idx
            break
    if anchor_index is None:
        raise TemplateUpdateError("Could not find 'Acceptance Criteria' section")

    for idx in range(anchor_index + 1, len(doc.paragraphs)):
        para = doc.paragraphs[idx]
        if para.text.strip():
            _set_paragraph_text(para, acceptance_text)
            return

    raise TemplateUpdateError("Could not find acceptance criteria value line")


def _iter_table_paragraphs(table) -> Iterable:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                yield p


def _iter_all_paragraphs(doc: Document) -> Iterable:
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        yield from _iter_table_paragraphs(table)

    for section in doc.sections:
        for container in (section.header, section.first_page_header, section.even_page_header):
            if container is None:
                continue
            for p in container.paragraphs:
                yield p
            for table in container.tables:
                yield from _iter_table_paragraphs(table)


def _replace_project_no_everywhere(doc: Document, project_no: str) -> None:
    project_no = (project_no or "").strip()
    if not project_no:
        return

    pattern_project_hash = re.compile(r"(Project#)\s*([A-Za-z0-9-]+)")
    pattern_project_spec = re.compile(r"(Project\s*Specification\s*:\s*)([A-Za-z0-9-]+)", flags=re.IGNORECASE)
    for para in _iter_all_paragraphs(doc):
        text = para.text
        if "Project#" not in text and "Project Specification" not in text:
            continue
        updated = pattern_project_hash.sub(lambda m: f"{m.group(1)}{project_no}", text, count=1)
        updated = pattern_project_spec.sub(lambda m: f"{m.group(1)}{project_no}", updated, count=1)
        if updated != text:
            _set_paragraph_text(para, updated)


def _extract_blip_ids_and_widths(paragraph) -> list[Tuple[str, Optional[float]]]:
    from lxml import etree

    ns = {
        "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
        "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    }
    items: list[Tuple[str, Optional[float]]] = []
    for run in paragraph.runs:
        root = etree.fromstring(run._element.xml.encode())
        blips = root.xpath(".//a:blip", namespaces=ns)
        extents = root.xpath(".//wp:extent", namespaces=ns)
        width_inches = None
        if extents:
            cx = extents[0].get("cx")
            if cx and cx.isdigit():
                width_inches = int(cx) / 914400.0
        for b in blips:
            rid = b.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
            if rid:
                items.append((rid, width_inches))
    return items


def _ensure_decision_rule_block(doc: Document, decision_rule_source_path: Optional[Path]) -> None:
    if _try_find_paragraph_by_prefix(doc, "Decision Rule:") is not None:
        return

    anchor = (
        _try_find_paragraph_by_prefix(doc, "Product Drawing:")
        or _try_find_paragraph_by_prefix(doc, "Part drawing:")
        or _try_find_paragraph_by_prefix(doc, "Procedure for monitoring:")
    )
    if anchor is None:
        return

    heading_text = "Decision Rule:"
    heading_font_name = "Verdana"
    heading_bold = True
    heading_size_pt = 12.0
    image_blob: Optional[bytes] = None
    image_width_inches = 6.0

    if decision_rule_source_path and decision_rule_source_path.exists():
        src_doc = Document(decision_rule_source_path)
        decision_para = _try_find_paragraph_by_prefix(src_doc, "Decision Rule:")
        if decision_para is not None:
            heading_text = decision_para.text.strip() or heading_text
            if decision_para.runs:
                run = next((r for r in decision_para.runs if r.text.strip()), decision_para.runs[0])
                if run.font.name:
                    heading_font_name = run.font.name
                if run.font.size:
                    heading_size_pt = run.font.size.pt
                if run.bold is not None:
                    heading_bold = run.bold

            start_idx = 0
            for idx, para in enumerate(src_doc.paragraphs):
                if para._p is decision_para._p:
                    start_idx = idx
                    break
            end_idx = len(src_doc.paragraphs)
            for idx in range(start_idx + 1, len(src_doc.paragraphs)):
                txt = src_doc.paragraphs[idx].text.strip()
                if txt.startswith("Part drawing:") or txt.startswith("Product Drawing:"):
                    end_idx = idx
                    break

            found = False
            for idx in range(start_idx, end_idx):
                para = src_doc.paragraphs[idx]
                for rid, width in _extract_blip_ids_and_widths(para):
                    related = src_doc.part.related_parts.get(rid)
                    if related is not None:
                        image_blob = related.blob
                        if width:
                            image_width_inches = width
                        found = True
                        break
                if found:
                    break

    p_decision = anchor.insert_paragraph_before(heading_text)
    if not p_decision.runs:
        p_decision.add_run(heading_text)
    for run in p_decision.runs:
        run.font.name = heading_font_name
        run.bold = heading_bold
        run.font.size = Pt(heading_size_pt)

    if image_blob:
        p_image = anchor.insert_paragraph_before("")
        p_image.add_run().add_picture(io.BytesIO(image_blob), width=Inches(image_width_inches))


def _try_find_paragraph_by_prefix(doc: Document, prefix: str, occurrence: int = 1):
    count = 0
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            count += 1
            if count == occurrence:
                return p
    return None


def _clear_paragraph_by_prefix(doc: Document, prefix: str, occurrence: int = 1) -> None:
    para = _try_find_paragraph_by_prefix(doc, prefix, occurrence=occurrence)
    if para is not None:
        _set_paragraph_text(para, "")


def _extract_first_number(text: str) -> Optional[float]:
    match = re.search(r"[-+]?[0-9]*\.?[0-9]+", text or "")
    if not match:
        return None
    try:
        return float(match.group(0))
    except ValueError:
        return None


def _format_hours_label(hours_value: float) -> str:
    if abs(hours_value - round(hours_value)) < 1e-6:
        return f"{int(round(hours_value))}"
    return f"{hours_value:.2f}".rstrip("0").rstrip(".")


def _testing_duration_label(num_samples: int, each_hours: float) -> str:
    total_hours = num_samples * each_hours
    if total_hours <= 0:
        return ""

    total_days = total_hours / 24.0
    if total_days <= 28:
        weeks = max(1, math.ceil(total_days / 7.0))
        return f"{weeks} week" if weeks == 1 else f"{weeks} weeks"

    months = max(1, math.ceil(total_days / 30.0))
    return f"{months} month" if months == 1 else f"{months} months"


def _find_on_test_end_row(ws: openpyxl.worksheet.worksheet.Worksheet, start_row: int = 8) -> int:
    # Prefer stopping before "Removal" section if present.
    for row in range(start_row + 1, ws.max_row + 1):
        marker = _safe_text(ws.cell(row, 1).value).lower()
        if marker == "removal":
            return row - 1

    # Fallback: last non-empty row in A:O.
    last = start_row
    for row in range(start_row, ws.max_row + 1):
        if any(_safe_text(ws.cell(row, col).value) for col in range(1, 16)):
            last = row
    return last


def _fill_first_empty_between(doc: Document, start_prefix: str, end_prefix: str, text: str):
    start_index: Optional[int] = None
    end_index: Optional[int] = None

    for idx, para in enumerate(doc.paragraphs):
        stripped = para.text.strip()
        if start_index is None and stripped.startswith(start_prefix):
            start_index = idx
        if stripped.startswith(end_prefix):
            end_index = idx
            if start_index is not None and end_index > start_index:
                break

    if start_index is None or end_index is None or end_index <= start_index:
        return None

    for idx in range(start_index + 1, end_index):
        para = doc.paragraphs[idx]
        if not para.text.strip():
            _set_paragraph_text(para, text)
            return para
    return None


def _fill_first_empty_between_any_end(doc: Document, start_prefix: str, end_prefixes: list[str], text: str):
    start_index: Optional[int] = None
    end_index: Optional[int] = None

    for idx, para in enumerate(doc.paragraphs):
        stripped = para.text.strip()
        if start_index is None and stripped.startswith(start_prefix):
            start_index = idx
            continue
        if start_index is None:
            continue
        if any(stripped.startswith(prefix) for prefix in end_prefixes):
            end_index = idx
            break

    if start_index is None:
        return None

    if end_index is None:
        end_index = len(doc.paragraphs)

    for idx in range(start_index + 1, end_index):
        para = doc.paragraphs[idx]
        if not para.text.strip():
            _set_paragraph_text(para, text)
            return para
    return None


def update_word_template(
    template_path: Path,
    output_path: Path,
    data: ExtractedData,
    pre_post_image_path: Path,
    duty_cycle_image_path: Path,
    report_date: Optional[str] = None,
    revision_no: Optional[str] = None,
    revision_date: Optional[str] = None,
    project_no: Optional[str] = None,
    project_leader: Optional[str] = None,
    tooling_lead_time: Optional[str] = None,
    decision_rule_source_path: Optional[Path] = None,
    product_drawing_image_path: Optional[Path] = None,
) -> None:
    doc = Document(template_path)

    part_compact = data.part_number.replace(" ", "")
    title = f"{part_compact} Seal  {data.purpose} with Oil fill".strip()

    if len(doc.tables) >= 2:
        title_table = doc.tables[1]
        _set_cell_text(title_table.rows[0].cells[1], title)
        _set_cell_text(title_table.rows[0].cells[2], title)
        _set_cell_text(title_table.rows[0].cells[3], title)
        _set_cell_text(title_table.rows[1].cells[1], data.requester)
        if project_leader is not None and len(title_table.rows[1].cells) >= 4:
            project_leader_text = project_leader.strip()
            if project_leader_text:
                _set_cell_text(title_table.rows[1].cells[3], project_leader_text)

    if len(doc.tables) >= 3:
        time_table = doc.tables[2]
        if len(time_table.rows) >= 2 and len(time_table.rows[1].cells) >= 2 and tooling_lead_time:
            _set_cell_text(time_table.rows[1].cells[1], tooling_lead_time)

        num_samples_num = int(_extract_first_number(data.num_samples) or 0)
        each_hours = _extract_first_number(data.acceptance_duration) or 0.0
        each_hours_label = _format_hours_label(each_hours) if each_hours > 0 else data.acceptance_duration.strip()
        if num_samples_num > 0:
            sample_word = "sample" if num_samples_num == 1 else "samples"
            testing_activity = f"Testing of {num_samples_num} {sample_word} @ {each_hours_label} Hrs. each"
            testing_duration = _testing_duration_label(num_samples_num, each_hours)
            if len(time_table.rows) >= 3 and len(time_table.rows[2].cells) >= 2:
                _set_cell_text(time_table.rows[2].cells[0], testing_activity)
                if testing_duration:
                    _set_cell_text(time_table.rows[2].cells[1], testing_duration)

    if len(doc.tables) >= 1:
        meta_table = doc.tables[0]
        if report_date and len(meta_table.rows[0].cells) >= 2:
            _set_cell_text(meta_table.rows[0].cells[1], report_date)
        if len(meta_table.rows[0].cells) >= 4:
            _set_cell_text(meta_table.rows[0].cells[3], revision_date or "DD/MM/YYYY")
        if revision_no is not None and len(meta_table.rows[0].cells) >= 6:
            revision_text = str(revision_no).strip()
            if revision_text:
                _set_cell_text(meta_table.rows[0].cells[5], revision_text)

    _replace_project_no_everywhere(doc, project_no or "")

    objective = (
        f"To Conduct the {data.purpose} with Oil fill on {data.num_samples} "
        f"samples of {part_compact} shaft seal."
    )
    _set_paragraph_text(_find_paragraph_by_prefix(doc, "To Conduct the"), objective)

    _set_paragraph_text(_find_paragraph_by_prefix(doc, "Shaft:"), f"Shaft: {data.shaft_material}")
    _set_paragraph_text(
        _find_paragraph_by_prefix(doc, "Diameter:", occurrence=1),
        f"Diameter: {data.shaft_diameter} ± {data.shaft_tolerance} {data.shaft_unit}",
    )
    _set_paragraph_text(
        _find_paragraph_by_prefix(doc, "Surface Finish:", occurrence=1),
        f"Surface Finish: {_normalize_ra(data.shaft_ra)}",
    )
    _set_paragraph_text(
        _find_paragraph_by_prefix(doc, "Hardness:"),
        f"Hardness: {data.shaft_hardness.replace('HRc', 'HRC')} Min.",
    )
    _set_paragraph_text(
        _find_paragraph_by_prefix(doc, "DRO:"),
        f"DRO: {data.dro_value} ± {data.dro_tolerance} {data.shaft_unit} TIR",
    )

    _set_paragraph_text(_find_paragraph_by_prefix(doc, "Bore:"), f"Bore: {data.bore_material}")
    _set_paragraph_text(
        _find_paragraph_by_prefix(doc, "Diameter:", occurrence=2),
        f"Diameter: {data.bore_diameter} ± {data.bore_tolerance}{data.bore_unit}",
    )
    _set_paragraph_text(
        _find_paragraph_by_prefix(doc, "STBM:"),
        f"STBM: {data.stbm_value} ± {data.stbm_tolerance} {data.bore_unit} TIR",
    )
    _set_paragraph_text(
        _find_paragraph_by_prefix(doc, "Surface Finish:", occurrence=2),
        f"Surface Finish: {_normalize_ra(data.bore_ra)}",
    )
    _set_paragraph_text(
        _find_paragraph_by_prefix(doc, "STBM Orientation:"),
        f"STBM Orientation: {data.contamination_stbm_orientation}",
    )
    _set_paragraph_text(
        _find_paragraph_by_prefix(doc, "Seals cock:"),
        f"Seals cock: {data.seal_cock_value} ± {data.seal_cock_tolerance} {data.shaft_unit}",
    )
    inserted_recip = _fill_first_empty_between(
        doc,
        "Seals cock:",
        "Fluid:",
        f"Reciprocation: {data.recip_value} ± {data.recip_tolerance} {data.shaft_unit}",
    )
    seals_cock_para = _try_find_paragraph_by_prefix(doc, "Seals cock:")
    if inserted_recip is not None and seals_cock_para is not None:
        _set_paragraph_font_like(inserted_recip, seals_cock_para)

    customer_caps = data.customer.upper() if data.customer else ""
    fluid_type = f"{data.oil_type} {customer_caps}".strip()
    _set_paragraph_text(_find_paragraph_by_prefix(doc, "Type –"), f"Type – {fluid_type}")
    _set_paragraph_text(
        _find_paragraph_by_prefix(doc, "Oil Change Interval:"),
        f"Oil Change Interval: {data.oil_change_interval}",
    )

    pre_lube_value = data.setup_notes or data.pre_lube_required
    _set_paragraph_text(_find_paragraph_by_prefix(doc, "Pre-lube:"), f"Pre-lube: {pre_lube_value}")

    setup_notes_text = (data.setup_notes_full or "").strip()
    if setup_notes_text:
        inserted_setup = _fill_first_empty_between_any_end(
            doc,
            "Pre-lube:",
            ["Contamination:", "Slurry Mixing:", "Test Procedure:"],
            f"Setup Notes: {setup_notes_text}",
        )
        fluid_ref = _try_find_paragraph_by_prefix(doc, "Fluid Level") or _try_find_paragraph_by_prefix(doc, "Type –")
        if inserted_setup is not None and fluid_ref is not None:
            _set_paragraph_font_like(inserted_setup, fluid_ref)

    contamination_content_present = any(
        x.strip()
        for x in [
            data.contamination_type,
            data.contamination_mix_ratio,
            data.contamination_amount,
            data.contamination_recip_freq,
        ]
    )

    if contamination_content_present and _is_contamination_test(data.test_type, data.purpose):
        heading_para = _try_find_paragraph_by_prefix(doc, "Slurry Mixing:")
        if heading_para is not None:
            _set_paragraph_text(heading_para, "Contamination:")
        _set_paragraph_text(_find_paragraph_by_prefix(doc, "Type:", occurrence=1), f"Type: {data.contamination_type}")
        _set_paragraph_text(_find_paragraph_by_prefix(doc, "Mix Ratio:"), f"Mix Ratio: {data.contamination_mix_ratio}")
        _set_paragraph_text(_find_paragraph_by_prefix(doc, "Amount:"), f"Amount: {data.contamination_amount}")
        amount_para = _try_find_paragraph_by_prefix(doc, "Amount:")
        inserted_recip_freq = _fill_first_empty_between(
            doc,
            "Amount:",
            "Test Procedure:",
            f"Recip. Frequency: {data.contamination_recip_freq}",
        )
        inserted_stbm_orient = _fill_first_empty_between(
            doc,
            "Amount:",
            "Test Procedure:",
            f"STBM Orientation: {data.contamination_stbm_orientation}",
        )
        if amount_para is not None:
            if inserted_recip_freq is not None:
                _set_paragraph_font_like(inserted_recip_freq, amount_para)
            if inserted_stbm_orient is not None:
                _set_paragraph_font_like(inserted_stbm_orient, amount_para)
    else:
        _clear_paragraph_by_prefix(doc, "Slurry Mixing:")
        _clear_paragraph_by_prefix(doc, "Contamination:")
        _clear_paragraph_by_prefix(doc, "Type:", occurrence=1)
        _clear_paragraph_by_prefix(doc, "Mix Ratio:")
        _clear_paragraph_by_prefix(doc, "Amount:")

    duration_clean = data.acceptance_duration.strip()
    _set_paragraph_text(_find_paragraph_by_prefix(doc, "The total test duration is"), f"The total test duration is {duration_clean}")
    _replace_acceptance_line(doc, data.acceptance_failure.strip())

    _replace_image_after_anchor(
        doc,
        "Pre and Post-test measurements of Seals are as follows",
        pre_post_image_path,
        width_inches=6.4,
    )
    _replace_image_after_anchor(doc, "Test cycle would be as follows:", duty_cycle_image_path, width_inches=7.0)
    if product_drawing_image_path and product_drawing_image_path.exists():
        if _try_find_paragraph_by_prefix(doc, "Product Drawing:") is not None:
            _replace_image_after_anchor(doc, "Product Drawing:", product_drawing_image_path, width_inches=6.6)
        elif _try_find_paragraph_by_prefix(doc, "Part drawing:") is not None:
            _replace_image_after_anchor(doc, "Part drawing:", product_drawing_image_path, width_inches=6.6)
    _ensure_decision_rule_block(doc, decision_rule_source_path)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def convert(
    excel_path: Path,
    template_docx_path: Path,
    output_docx_path: Path,
    temp_dir: Path,
    report_date: Optional[str] = None,
    revision_no: Optional[str] = None,
    revision_date: Optional[str] = None,
    project_no: Optional[str] = None,
    project_leader: Optional[str] = None,
    tooling_lead_time: Optional[str] = None,
    decision_rule_source_path: Optional[Path] = None,
) -> None:
    data = extract_excel_data(excel_path)

    wb = openpyxl.load_workbook(excel_path, data_only=True)
    page1 = wb["Page 1"]
    page2 = wb["Page 2"]

    pre_post_img = temp_dir / "pre_post_measurements.png"
    duty_cycle_img = temp_dir / "duty_cycle.png"
    product_drawing_img = temp_dir / "product_drawing_main.png"

    render_sheet_range_to_image(page1, "A30:L44", pre_post_img, end_at_last_content=False)
    on_test_end_row = _find_on_test_end_row(page2, start_row=8)
    render_sheet_range_to_image(page2, f"A8:O{on_test_end_row}", duty_cycle_img, end_at_last_content=True)
    product_drawing_image_path = _extract_main_drawing_image(wb, product_drawing_img, sheet_name="Part Drawing 1")

    if decision_rule_source_path is None:
        candidate = Path(__file__).resolve().parents[1] / "assets" / "Project Specification - Decision Rule Source.docx"
        if candidate.exists():
            decision_rule_source_path = candidate

    update_word_template(
        template_docx_path,
        output_docx_path,
        data,
        pre_post_img,
        duty_cycle_img,
        report_date=report_date,
        revision_no=revision_no,
        revision_date=revision_date,
        project_no=project_no,
        project_leader=project_leader,
        tooling_lead_time=tooling_lead_time,
        decision_rule_source_path=decision_rule_source_path,
        product_drawing_image_path=product_drawing_image_path,
    )


def _parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Convert SKF request Excel sheet into Word project specification based on a template.",
    )
    parser.add_argument("--excel", required=True, type=Path, help="Path to source .xlsm file")
    parser.add_argument("--template", required=True, type=Path, help="Path to template .docx file")
    parser.add_argument("--output", required=True, type=Path, help="Path for generated .docx")
    parser.add_argument(
        "--temp-dir",
        type=Path,
        default=Path("output") / "generated_assets",
        help="Folder for intermediate generated images",
    )
    parser.add_argument("--date", dest="report_date", type=str, default=None, help="Original date (DD/MM/YYYY)")
    parser.add_argument("--revision", dest="revision_no", type=str, default=None, help="Revision number (0/1/2...)")
    parser.add_argument(
        "--revision-date",
        dest="revision_date",
        type=str,
        default=None,
        help="Revision date (DD/MM/YYYY). If not set, DD/MM/YYYY is used in report.",
    )
    parser.add_argument(
        "--project-no",
        dest="project_no",
        type=str,
        default=None,
        help="Project number to place next to Project# in report headers.",
    )
    parser.add_argument(
        "--project-leader",
        dest="project_leader",
        type=str,
        default=None,
        help="Project leader name to place in report table.",
    )
    parser.add_argument(
        "--tooling-lead-time",
        dest="tooling_lead_time",
        type=str,
        default=None,
        help="Estimated lead time for Tooling Design, Manufacture and Inspection.",
    )
    parser.add_argument(
        "--decision-rule-source",
        dest="decision_rule_source",
        type=Path,
        default=None,
        help="Optional source .docx to copy Decision Rule heading/image block from.",
    )
    return parser.parse_args()


def main() -> None:
    args = _parse_args()
    convert(
        args.excel,
        args.template,
        args.output,
        args.temp_dir,
        report_date=args.report_date,
        revision_no=args.revision_no,
        revision_date=args.revision_date,
        project_no=args.project_no,
        project_leader=args.project_leader,
        tooling_lead_time=args.tooling_lead_time,
        decision_rule_source_path=args.decision_rule_source,
    )
    print(f"Generated Word document: {args.output}")


if __name__ == "__main__":
    main()
