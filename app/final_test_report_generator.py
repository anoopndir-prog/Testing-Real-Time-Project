"""Generate the fixed-format Final Test Report from user inputs.

The generator starts from a Word template so page setup, headers, footers,
fonts, borders, and table geometry stay as close as possible to the approved
format. It then replaces only the extracted/user-supplied report values.
"""

from __future__ import annotations

import copy
import datetime as dt
import io
import re
import shutil
import subprocess
import tempfile
import zipfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
import openpyxl

from app.paths import resource_path
from app.report_export import convert_docx_to_pdf as _convert_docx_to_pdf


TEMPLATE_PATH = resource_path(Path("assets/Final Test Report - Template.docx"))
PROJECT_NO_RE = re.compile(r"\bTR\d{2}-\d{4}-BTS\b", re.IGNORECASE)
RESULT_OPTIONS = ["PASS", "FAIL", "No results defined", "No acceptance Criteria defined"]


@dataclass
class DistributionEntry:
    name: str
    organisation: str
    location: str


@dataclass
class SamplePhotoEntry:
    """One sample's Section 5 Before/After Test photos. `before_caption`
    is free text the user supplies (what the photo shows, e.g. "Shaft,
    Bore") since that varies photo to photo and can't be inferred; the
    After-Test photo's caption isn't user-supplied at all - it's filled
    in automatically from that sample's own remarks (the same text
    Section 2's Conclusions table shows, e.g. "Zero Leakage observed")."""

    before_photo: bytes | None = None
    before_caption: str = ""
    after_photo: bytes | None = None


@dataclass
class FinalReportInputs:
    project_spec_path: Path
    inspection_sheet_path: Path
    monitoring_sheet_paths: list[Path]
    report_date: str
    project_leader: str
    reviewer: str
    approved_by: str
    reviewer_role: str
    approved_by_role: str
    distribution: list[DistributionEntry]
    # Already-cropped seal identification photo (JPEG bytes) for section
    # 3.2. Optional: when absent, the template's example photo is left
    # in place rather than the report shipping with a blank slot.
    seal_photo: bytes | None = None
    # Section 5's per-sample Before/After Test photos (JPEG bytes, already
    # cropped), keyed by the same "<part_no> Sample #NN" label
    # `detect_sample_labels`/`_sample_photo_label` produce - lets the
    # caller collect photos before generation using labels detected from
    # just the Project Spec + Inspection sheet, independent of exactly
    # when generation itself recomputes the sample list. Samples with no
    # matching entry are simply left blank.
    sample_photos: dict[str, SamplePhotoEntry] | None = None


@dataclass
class TestSpecValue:
    """One row of the Test description table below Direction of Rotation:
    a single Specification-as-per-OSTR/Drawing value plus, where the value
    can vary sample to sample (measured dimensions), the actual reading
    keyed by sample suffix (e.g. "S#01")."""

    spec: str = ""
    by_sample: dict[str, str] | None = None

    def value_for(self, sample: str, fallback: str = "") -> str:
        if self.by_sample:
            value = self.by_sample.get(_sample_suffix(sample))
            if value:
                return value
        return fallback


@dataclass
class ExtractedFinalReportData:
    project_no: str = ""
    project_title: str = ""
    objective: str = ""
    requester_name: str = ""
    requester_organisation: str = ""
    requester_location: str = ""
    target_hours: str = ""
    actual_hours: str = ""
    test_request_form_no: str = ""
    rotation_direction: str = ""
    sample_receipt_date: str = ""
    seal_part_no: str = ""
    requested_sample_count: int = 0
    tested_sample_count: int = 0
    sample_numbers: list[str] | None = None
    sample_results: list["SampleResult"] | None = None
    remarks: list[str] | None = None
    suspension_note: str = ""
    section_b_measurements: dict[str, dict[str, tuple[str, str]]] | None = None
    section_b_specs: dict[str, str] | None = None
    section_b_remarks: str = ""
    measured_by: str = ""
    checked_by: str = ""
    dro: TestSpecValue = field(default_factory=TestSpecValue)
    stbm: TestSpecValue = field(default_factory=TestSpecValue)
    seal_cock: TestSpecValue = field(default_factory=TestSpecValue)
    reciprocation: TestSpecValue = field(default_factory=TestSpecValue)
    shaft_surface_finish: TestSpecValue = field(default_factory=TestSpecValue)
    shaft_diameter: TestSpecValue = field(default_factory=TestSpecValue)
    shaft_hardness: TestSpecValue = field(default_factory=TestSpecValue)
    housing_diameter: TestSpecValue = field(default_factory=TestSpecValue)
    housing_surface_finish: TestSpecValue = field(default_factory=TestSpecValue)
    fluid: TestSpecValue = field(default_factory=TestSpecValue)
    fluid_level: TestSpecValue = field(default_factory=TestSpecValue)
    oil_change_interval: TestSpecValue = field(default_factory=TestSpecValue)
    raw_project_text: str = ""
    test_dates: list["TestDateEntry"] | None = None


@dataclass
class TestDateEntry:
    sample_no: str
    start_date: str
    end_date: str
    position: str
    machine_identification: str


@dataclass
class SampleResult:
    sample_no: str
    target_hours: str
    actual_hours: str
    result: str
    remarks: str


def generate_final_report(
    inputs: FinalReportInputs, output_format: str, output_dir: Path | None = None
) -> Path:
    """Build the report and return the path it was written to.

    `output_dir` defaults to the desktop app's historical behavior (the
    user's own Downloads folder). The web app passes a temp directory
    instead, since a shared server must never write into its own host
    machine's Downloads folder on a caller's behalf."""
    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(f"Final Report template not found: {TEMPLATE_PATH}")

    extracted = _extract_data(
        inputs.project_spec_path,
        inputs.inspection_sheet_path,
        inputs.monitoring_sheet_paths,
    )
    if inputs.project_leader.strip():
        extracted.measured_by = inputs.project_leader.strip()

    project_no = extracted.project_no or _project_no_from_filename(inputs.project_spec_path)
    if not project_no:
        project_no = "TRXX-XXXX-BTS"

    title = extracted.project_title or f"Final Test Report - {project_no}"
    report_date = _display_date(inputs.report_date)
    safe_name = _safe_filename(f"Final report - {project_no}")
    target_dir = output_dir if output_dir is not None else Path.home() / "Downloads"
    target_dir.mkdir(parents=True, exist_ok=True)
    out_docx = target_dir / f"{safe_name}.docx"

    shutil.copyfile(TEMPLATE_PATH, out_docx)
    doc = Document(out_docx)

    author = inputs.project_leader.strip()
    reviewer = inputs.reviewer.strip()
    approver = inputs.approved_by.strip()

    _replace_known_text(doc, {
        "TR25-0009-BTS": project_no,
        "01 April 2025": report_date,
        "71261 – ID 616493 - Seal Endurance Test with Oil Fill": title,
        "Endurance test for Pinion Application – VECV": title,
        "Anoop N": author,
        "Dhayalan K": approver,
    })

    _set_report_data(doc, project_no, report_date, title)
    _set_people_block(doc, author, reviewer, approver)
    _set_distribution(doc, inputs, extracted)
    _set_project_team(doc, inputs, extracted)
    _set_objective(doc, extracted.objective)
    _set_conclusion_table(doc, extracted)
    _set_final_results_statement(doc, extracted)
    _set_sample_receipt(doc, extracted)
    _set_sample_identification(doc, extracted)
    if inputs.seal_photo:
        _set_seal_identification_photo(doc, inputs.seal_photo)
    _set_photo_placeholders(doc, extracted, inputs.sample_photos)
    _set_test_description_table(doc, extracted)
    _set_procedure_and_criteria_text(doc, extracted.raw_project_text)
    _set_test_dates_table(doc, extracted)
    _set_pre_post_measurement_tables(doc, extracted)
    _set_footer_project_no(doc, project_no)
    _refresh_toc_page_numbers(doc)

    doc.save(out_docx)
    _patch_textbox_title(out_docx, title, author, project_no, report_date, approver)

    if output_format.lower() == "word":
        return out_docx

    out_pdf = out_docx.with_suffix(".pdf")
    _convert_docx_to_pdf(out_docx, out_pdf)
    return out_pdf


def _extract_data(
    project_spec_path: Path,
    inspection_sheet_path: Path,
    monitoring_sheet_paths: list[Path],
) -> ExtractedFinalReportData:
    project_text = _extract_text(project_spec_path)
    inspection_text = _extract_text(inspection_sheet_path)
    monitoring_texts = [_extract_text(path) for path in monitoring_sheet_paths]
    monitoring_text = "\n".join(monitoring_texts)
    combined = f"{project_text}\n{inspection_text}\n{monitoring_text}"

    data = ExtractedFinalReportData()
    data.project_no = _first_match(PROJECT_NO_RE, combined)
    data.project_title = _extract_project_title(project_text, data.project_no)
    data.objective = _extract_section_line(project_text, "Objective") or _extract_objective_sentence(project_text)
    data.requester_name = _extract_requester_name(project_text)
    data.requester_organisation = _extract_label_value(project_text, ["Customer"])
    data.requester_location = "Bangalore, India"
    data.target_hours = _extract_target_hours(project_text)
    data.actual_hours = _extract_actual_hours(monitoring_text or inspection_text, data.target_hours)
    data.test_request_form_no = _extract_test_request_form_no(data.objective, project_text, data.project_no)
    data.rotation_direction = _extract_rotation_direction(project_text)
    data.sample_receipt_date = _extract_sample_receipt_date(inspection_text)
    data.seal_part_no = _extract_seal_part_no(project_text, data.project_title, project_spec_path.name)
    data.requested_sample_count = _extract_sample_count_from_text(project_text)
    data.tested_sample_count = _extract_tested_sample_count(inspection_text)
    data.sample_numbers = _select_tested_sample_numbers(
        combined,
        data.seal_part_no,
        data.requested_sample_count,
        data.tested_sample_count,
    )
    data.remarks = _extract_remarks(inspection_text)
    data.sample_results = _build_sample_results(data, inspection_text, monitoring_texts)
    data.suspension_note = _extract_suspension_note(combined)
    seal_used_map = _combined_seal_used_map(inspection_text)
    data.section_b_measurements, data.section_b_specs, data.section_b_remarks = _extract_section_b_measurements(inspection_text, seal_used_map)
    data.measured_by = _extract_label_value(inspection_text, ["Measured by", "Measured By"])
    data.checked_by = _extract_label_value(inspection_text, ["Checked by", "Checked By"])
    _populate_test_spec_values(data, project_text, inspection_text)
    data.raw_project_text = project_text
    data.test_dates = _build_test_dates(data, inspection_text)
    return data


def _populate_test_spec_values(data: ExtractedFinalReportData, project_text: str, inspection_text: str) -> None:
    """Fill in the Test description rows below Direction of Rotation:
    the Specification-as-per-OSTR/Drawing value from the Project
    Specification's Test Specification section, plus the actual
    per-sample readings handwritten into the Inspection & Execution
    sheet (Section C measurements, and the Tooling Measurement table
    mapped to samples via the "Seal used" rows)."""
    shaft_text = _slice_between(project_text, "Shaft", ["Bore", "Fluid", "Test Procedure"])
    bore_text = _slice_between(project_text, "Bore", ["Fluid", "Test Procedure"])
    fluid_text = _slice_between(project_text, "Fluid", ["Test Procedure", "Product Drawing", "Acceptance Criteria"])

    section_c = _extract_section_c_rows(inspection_text)
    data.dro = TestSpecValue(
        spec=_extract_dashed_value(shaft_text, ["DRO"]),
        by_sample={sample: values.get("DRO", "") for sample, values in section_c.items() if values.get("DRO")},
    )
    data.stbm = TestSpecValue(
        spec=_extract_dashed_value(bore_text, ["STBM"]),
        by_sample={sample: values.get("STBM", "") for sample, values in section_c.items() if values.get("STBM")},
    )
    data.seal_cock = TestSpecValue(
        spec=_extract_dashed_value(bore_text, ["Seal cock", "Seal Cock"]),
        by_sample={sample: values.get("Seal Cock", "") for sample, values in section_c.items() if values.get("Seal Cock")},
    )
    # Reciprocation is only asked for by some requestors; only populate the
    # row at all when the project specification actually mentions it.
    reciprocation_spec = _extract_dashed_value(shaft_text, ["Reciprocation"]) or _extract_dashed_value(bore_text, ["Reciprocation"])
    if reciprocation_spec:
        data.reciprocation = TestSpecValue(
            spec=reciprocation_spec,
            by_sample={sample: values.get("Reciprocation", "") for sample, values in section_c.items() if values.get("Reciprocation")},
        )

    tooling = _extract_tooling_measurements(inspection_text)
    shaft_map = _extract_seal_used_mapping(inspection_text, "shaft")
    bore_map = _extract_seal_used_mapping(inspection_text, "bore")

    data.shaft_surface_finish = TestSpecValue(
        spec=_extract_dashed_value(shaft_text, ["Surface Roughness", "Surface Finish", "Roughness"]),
        by_sample=_values_by_sample(shaft_map, tooling.get("Roughness, Ra (µm)", {})),
    )
    data.shaft_diameter = TestSpecValue(
        spec=_extract_dashed_value(shaft_text, ["Diameter"]),
        by_sample=_values_by_sample(shaft_map, tooling.get("OD (mm)", {})),
    )
    data.shaft_hardness = TestSpecValue(
        spec=_extract_dashed_value(shaft_text, ["Hardness"]),
        by_sample=_values_by_sample(shaft_map, tooling.get("Hardness (HRC)", {})),
    )
    data.housing_diameter = TestSpecValue(
        spec=_extract_dashed_value(bore_text, ["Diameter"]),
        by_sample=_values_by_sample(bore_map, tooling.get("ID (mm)", {})),
    )
    data.housing_surface_finish = TestSpecValue(
        spec=_extract_dashed_value(bore_text, ["Surface Roughness", "Surface Finish", "Roughness"]),
        by_sample=_values_by_sample(bore_map, tooling.get("Roughness, Ra (µm) [Bore]", {})),
    )

    data.fluid = TestSpecValue(spec=_extract_dashed_value(fluid_text, ["Oil type", "Fluid type", "Type"]))
    data.fluid_level = TestSpecValue(spec=_extract_dashed_value(fluid_text, ["Fluid Level"]))
    data.oil_change_interval = TestSpecValue(spec=_extract_dashed_value(fluid_text, ["Oil Change Interval"]))


def _extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return _extract_docx_text(path)
    if suffix == ".pdf":
        return _extract_pdf_text(path)
    if suffix in {".xlsx", ".xlsm"}:
        return _extract_excel_text(path)
    return ""


def _extract_docx_text(path: Path) -> str:
    doc = Document(path)
    parts: list[str] = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text.strip())
    for table in doc.tables:
        for row in table.rows:
            vals = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if vals:
                parts.append(" | ".join(vals))
    return "\n".join(parts)


def _extract_pdf_text(path: Path) -> str:
    for module_name in ("pypdf", "PyPDF2"):
        try:
            module = __import__(module_name)
            reader = module.PdfReader(str(path))
            text = "\n".join((page.extract_text() or "") for page in reader.pages)
        except Exception:
            text = ""
        if text.strip():
            return text

    try:
        result = subprocess.run(
            ["pdftotext", str(path), "-"],
            check=True,
            capture_output=True,
            text=True,
            timeout=30,
        )
        if result.stdout.strip():
            return result.stdout
    except Exception:
        pass

    # Daily monitoring sheets are frequently a scan of a handwritten log
    # rather than a digital PDF, so they carry no extractable text layer
    # at all. Fall back to OCR in that case.
    return _extract_pdf_text_via_ocr(path)


def _extract_pdf_text_via_ocr(path: Path) -> str:
    """Best-effort OCR for scanned/handwritten PDFs. Requires the PyMuPDF
    and pytesseract packages plus a local Tesseract OCR install; returns
    "" silently if any of those are missing, so callers must not assume
    OCR text is available or fully accurate for handwritten entries."""
    try:
        import fitz  # PyMuPDF
        import pytesseract
        from PIL import Image
    except Exception:
        return ""

    try:
        doc = fitz.open(str(path))
    except Exception:
        return ""

    parts: list[str] = []
    try:
        for page in doc:
            pix = page.get_pixmap(matrix=fitz.Matrix(3, 3))
            image = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
            try:
                parts.append(pytesseract.image_to_string(image))
            except Exception:
                continue
    finally:
        doc.close()
    return "\n".join(parts)


def _extract_excel_text(path: Path) -> str:
    try:
        wb = openpyxl.load_workbook(path, data_only=True, read_only=True)
    except Exception:
        return ""

    parts: list[str] = []
    for ws in wb.worksheets:
        parts.append(ws.title)
        for row in ws.iter_rows(values_only=True):
            values = [str(value).strip() for value in row if value not in (None, "")]
            if values:
                parts.append(" | ".join(values))
    return "\n".join(parts)


def _replace_known_text(doc: Document, replacements: dict[str, str]) -> None:
    for paragraph in _iter_paragraphs(doc):
        _replace_in_paragraph(paragraph, replacements)


def _replace_in_paragraph(paragraph, replacements: dict[str, str]) -> None:
    if not paragraph.runs:
        return
    original = paragraph.text
    updated = original
    for old, new in replacements.items():
        if new:
            updated = updated.replace(old, new)
    if updated == original:
        return
    for run in paragraph.runs:
        run.text = ""
    paragraph.runs[0].text = updated


def _iter_paragraphs(doc: Document) -> Iterable:
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p
    for section in doc.sections:
        for container in (
            section.header,
            section.first_page_header,
            section.even_page_header,
            section.footer,
            section.first_page_footer,
            section.even_page_footer,
        ):
            for p in container.paragraphs:
                yield p
            for table in container.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            yield p


def _set_report_data(doc: Document, project_no: str, report_date: str, title: str) -> None:
    labels = {
        "Report number": project_no,
        "Unit": "Global Testing India",
        "Location": "Bangalore",
        "Report type": "Final report",
        "Publication date": report_date,
        "Request number": project_no,
        "Customer project number": "NA",
        "Classification": "CONFIDENTIAL",
        "Keywords": title,
        "Related reports": "NA",
    }
    for idx, p in enumerate(doc.paragraphs):
        text = p.text.strip().replace("\t", "")
        if text in labels and idx + 1 < len(doc.paragraphs):
            _set_paragraph_text(doc.paragraphs[idx + 1], labels[text])


def _set_people_block(doc: Document, author: str, reviewer: str, approver: str) -> None:
    values = {"Author": author, "Reviewer": reviewer, "Approved by": approver}
    for idx, p in enumerate(doc.paragraphs):
        key = p.text.strip().replace("\t", "")
        if key in values and idx + 1 < len(doc.paragraphs):
            _set_paragraph_text(doc.paragraphs[idx + 1], values[key])


def _set_distribution(doc: Document, inputs: FinalReportInputs, data: ExtractedFinalReportData) -> None:
    if len(doc.tables) < 2:
        return

    requester_table = doc.tables[0]
    _ensure_rows(requester_table, 2)
    _set_row_values(requester_table.rows[1], [
        data.requester_name or "Requester",
        data.requester_organisation or "Global Testing India",
        data.requester_location or "Bangalore, India",
    ])

    dist_table = doc.tables[1]
    entries = [e for e in inputs.distribution if any((e.name, e.organisation, e.location))]
    _resize_table(dist_table, max(2, len(entries) + 1))
    for row_idx in range(1, len(dist_table.rows)):
        entry = entries[row_idx - 1] if row_idx - 1 < len(entries) else DistributionEntry("", "", "")
        _set_row_values(dist_table.rows[row_idx], [entry.name, entry.organisation, entry.location])


def _set_project_team(doc: Document, inputs: FinalReportInputs, data: ExtractedFinalReportData) -> None:
    if len(doc.tables) < 3:
        return

    people: list[tuple[str, str, str]] = []
    _add_person(people, inputs.approved_by, inputs.approved_by_role, "Bangalore, India")
    _add_person(people, inputs.reviewer, inputs.reviewer_role, "Bangalore, India")
    _add_person(people, inputs.project_leader, "Lead engineer Testing - Global Testing India", "Bangalore, India")
    _add_person(people, data.measured_by, "Associate Technician Testing - Global Testing India", "Bangalore, India")
    _add_person(people, data.checked_by, "Technician Testing - Global Testing India", "Bangalore, India")

    table = doc.tables[2]
    _resize_table(table, max(2, len(people) + 1))
    for idx, row in enumerate(table.rows[1:], start=0):
        name, role, location = people[idx] if idx < len(people) else ("", "", "")
        _set_row_values(row, [name, role, location])


def _set_objective(doc: Document, objective: str) -> None:
    if not objective:
        return
    idx = _paragraph_index(doc, "Objective", start=45)
    if idx is not None and idx + 1 < len(doc.paragraphs):
        _set_paragraph_text(doc.paragraphs[idx + 1], objective)


_DOCX_SECTION_BOUNDARY_TEXTS = {
    "test procedure:", "product drawing:", "acceptance criteria:",
    "procedure for monitoring:", "disclaimer:", "decision rule:",
    "test and measuring equipment.", "date the test was performed.",
}


def _set_procedure_and_criteria_text(doc: Document, project_text: str) -> None:
    """Copy the Test Procedure notes, Acceptance Criteria, Decision rule
    (if the spec has one), Procedure for monitoring, and Disclaimer
    verbatim from the attached Project Specification. These vary project
    to project (different tolerances, leak criteria, monitoring
    schedules), so the template's own boilerplate must not be left in a
    report for a different project."""
    # The paragraphs between "Test Procedure:" and "Product Drawing:" are
    # fragmented by two embedded images (the M1-M21 measurement table and
    # the Test Cycle table), but _content_paragraph_indices_after already
    # skips image-only (blank-text) paragraphs, so bullet-for-bullet
    # replacement works the same way as any other section here.
    _replace_section_paragraphs(doc, "Test Procedure:", _extract_bullet_lines(project_text, "Test Procedure:"))
    _replace_section_paragraphs(doc, "Acceptance Criteria:", _extract_bullet_lines(project_text, "Acceptance Criteria:"))
    _replace_section_paragraphs(doc, "Decision rule:", _extract_bullet_lines(project_text, "Decision rule:"))
    _replace_section_paragraphs(doc, "Procedure for monitoring:", _extract_bullet_lines(project_text, "Procedure for monitoring:"))
    _replace_section_paragraphs(doc, "Disclaimer:", _extract_bullet_lines(project_text, "Disclaimer:"))


def _replace_section_paragraphs(doc: Document, heading: str, lines: list[str]) -> None:
    """Replace the placeholder paragraph(s) under `heading` one-for-one
    with `lines`. If the source has fewer lines than placeholders, the
    extra placeholders are cleared; if it has more, the overflow is
    merged into the last placeholder rather than inserting new
    paragraphs (which would need to clone Word formatting blindly)."""
    if not lines:
        return
    heading_idx = _paragraph_index(doc, heading)
    if heading_idx is None:
        return
    content_indices = _content_paragraph_indices_after(doc, heading_idx)
    if not content_indices:
        return
    if len(lines) > len(content_indices):
        head = lines[:len(content_indices) - 1] if len(content_indices) > 1 else []
        tail = " ".join(lines[len(content_indices) - 1:])
        lines = [*head, tail] if head else [tail]
    for i, idx in enumerate(content_indices):
        _set_paragraph_text(doc.paragraphs[idx], lines[i] if i < len(lines) else "")


def _content_paragraph_indices_after(doc: Document, heading_idx: int, max_scan: int = 40) -> list[int]:
    """Collect the bullet/content paragraph indices under a heading. Blank
    spacer paragraphs are skipped rather than treated as the end of the
    section - the template has one between the 2nd and 3rd "Procedure for
    monitoring" bullets, for example - so only a known next-heading text
    (or the scan limit) ends the run."""
    indices: list[int] = []
    for idx in range(heading_idx + 1, min(heading_idx + 1 + max_scan, len(doc.paragraphs))):
        text = doc.paragraphs[idx].text.strip()
        if not text:
            continue
        if text.casefold() in _DOCX_SECTION_BOUNDARY_TEXTS:
            break
        indices.append(idx)
    return indices


def _effective_sample_results(data: ExtractedFinalReportData) -> list[SampleResult]:
    """The per-sample rows this report actually has results for - i.e. what
    Section 2's Conclusions table shows. Falls back to deriving a row per
    known sample number from the flat remarks/hours fields when the fuller
    per-sample extraction in _build_sample_results came up empty."""
    sample_results = data.sample_results or []
    if sample_results:
        return sample_results
    samples = data.sample_numbers or ["S#01"]
    remarks = data.remarks or []
    sample_results = []
    for idx, sample in enumerate(samples):
        remark = remarks[idx] if idx < len(remarks) else ""
        actual_hours = data.actual_hours
        result = _derive_result(data.target_hours, actual_hours, remark)
        clean_remark = _clean_result_from_remark(remark)
        sample_results.append(SampleResult(sample, data.target_hours, actual_hours, result, clean_remark))
    return sample_results


def _set_conclusion_table(doc: Document, data: ExtractedFinalReportData) -> None:
    if len(doc.tables) < 4:
        return
    sample_results = _effective_sample_results(data)
    table = doc.tables[3]
    _resize_table(table, len(sample_results) + 1)
    for idx, sample_result in enumerate(sample_results, start=1):
        row = table.rows[idx]
        values = [
            str(idx),
            sample_result.sample_no,
            sample_result.target_hours,
            sample_result.actual_hours,
            sample_result.result,
            sample_result.remarks,
        ]
        for col_idx, value in enumerate(values):
            if col_idx < len(row.cells):
                if col_idx == 4:
                    _set_cell_dropdown(row.cells[col_idx], value, RESULT_OPTIONS)
                else:
                    row.cells[col_idx].text = value


_NUMBER_WORDS = {1: "One", 2: "Two", 3: "Three", 4: "Four", 5: "Five", 6: "Six", 7: "Seven", 8: "Eight", 9: "Nine", 10: "Ten"}


def _number_word(count: int) -> str:
    return _NUMBER_WORDS.get(count, str(count))


def _set_final_results_statement(doc: Document, data: ExtractedFinalReportData) -> None:
    """Rewrite the templated pass/fail summary under "5 Test results" so it
    states how many samples this report actually covers and what happened
    to each - rather than leaving the template's own worked example ("3
    samples were run... zero Leakage... PASS the test") in a report where a
    different number of samples ran, or one of them leaked/failed. Driven
    by the same per-sample results as the Conclusions table (Section 2) and
    the test dates (Section 4.3), so the two sections can't disagree. The
    fixed "Important: Unless otherwise indicated..." notice right above it
    is boilerplate and is left untouched, as instructed."""
    heading_idx = _paragraph_index(doc, "Test results", start=150)
    if heading_idx is None or heading_idx + 2 >= len(doc.paragraphs):
        return
    statement = _build_final_results_statement(data)
    if statement:
        _set_paragraph_text(doc.paragraphs[heading_idx + 2], statement)


def _build_final_results_statement(data: ExtractedFinalReportData) -> str:
    results = _effective_sample_results(data)
    if not results:
        return ""

    count = len(results)
    sample_word = "sample" if count == 1 else "samples"
    sample_list = ", ".join(result.sample_no for result in results)
    leak_notes = list(dict.fromkeys(result.remarks.strip() for result in results if result.remarks and result.remarks.strip()))
    leak_clause = f", {'; '.join(leak_notes)}" if leak_notes else ""

    outcomes = {result.result for result in results}
    if outcomes == {"PASS"}:
        verdict = "PASS"
    elif "FAIL" in outcomes:
        verdict = "FAIL"
    else:
        verdict = ""

    subject = "the sample" if count == 1 else "all the samples"
    be_verb = "was" if count == 1 else "were"
    have_verb = "is" if count == 1 else "are"

    if verdict:
        verdict_clause = f"Hence, {subject} {have_verb} considered to have {verdict} the test."
    else:
        verdict_clause = f"The result for {subject} is not yet conclusive."

    sentence = (
        f"{_number_word(count)} {sample_word} ({sample_list}) {be_verb} run as per the test cycle"
        f"{leak_clause}. {verdict_clause}"
    )

    remaining = max((data.requested_sample_count or 0) - count, 0)
    if remaining:
        note = data.suspension_note
        if not note:
            remaining_word = _number_word(remaining)
            remaining_word_sample = "sample" if remaining == 1 else "samples"
            note = f"Testing on the remaining {remaining_word.lower()} {remaining_word_sample} was not completed."
        sentence = f"{sentence} {note}"

    return sentence


def _set_sample_receipt(doc: Document, data: ExtractedFinalReportData) -> None:
    idx = _paragraph_index(doc, "Date of receipt of the samples", start=55)
    if idx is None or idx + 1 >= len(doc.paragraphs):
        return
    sample_count = len(data.sample_results or data.sample_numbers or []) or _extract_sample_count_from_text(data.objective) or 1
    count_text = f"{sample_count:02d}" if sample_count < 10 else str(sample_count)
    receipt_date = data.sample_receipt_date or "DD.MM.YYYY"
    _set_paragraph_text(doc.paragraphs[idx + 1], f"{count_text} Samples received on {receipt_date}.")


def _set_sample_identification(doc: Document, data: ExtractedFinalReportData) -> None:
    if len(doc.tables) < 5:
        return
    samples = data.sample_numbers or [result.sample_no for result in data.sample_results or []] or ["S#01"]
    part_no = data.seal_part_no
    sample_labels = [_sample_suffix(sample) for sample in samples]
    sample_text = ", ".join(label for label in sample_labels if label)
    if part_no and sample_text:
        sample_text = f"{part_no} {sample_text}"
    table = doc.tables[4]
    _resize_table(table, 2)
    _set_row_values(table.rows[1], [sample_text, "SKF Sealing Solution Mysore"])


def _set_seal_identification_photo(doc: Document, photo_jpeg: bytes) -> None:
    """Drops the user-supplied (already auto-cropped) seal photo into
    section 3.2, replacing the template's example photo and its stray
    caption textbox - both are `w:drawing` elements living in the
    paragraphs between the "Identification of the samples" heading and
    the next heading."""
    heading_idx = _paragraph_index(doc, "Identification of the samples", start=55)
    if heading_idx is None:
        return
    end_idx = _paragraph_index(doc, "Description and condition of the test samples", start=heading_idx)
    if end_idx is None:
        end_idx = min(heading_idx + 15, len(doc.paragraphs))

    for idx in range(heading_idx, end_idx):
        for drawing in doc.paragraphs[idx]._p.xpath(".//w:drawing"):
            drawing.getparent().remove(drawing)

    target_idx = heading_idx + 1 if heading_idx + 1 < end_idx else heading_idx
    paragraph = doc.paragraphs[target_idx]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    picture = run.add_picture(io.BytesIO(photo_jpeg), width=Inches(3.5))
    _apply_picture_frame(picture, SEAL_PHOTO_BORDER_COLOR)


def _set_photo_placeholders(
    doc: Document,
    data: ExtractedFinalReportData,
    sample_photos: dict[str, "SamplePhotoEntry"] | None = None,
) -> None:
    """Fills the "Sample # | Photos - Before Test | Photos - After Test"
    table under Section 5 Test results. The sample list comes from
    `_effective_sample_results` - the same source Section 2's Conclusions
    table is built from, and consistent with Section 4.3's Date the test
    was performed table (`_build_test_dates`, driven by the same
    underlying sample numbers) - so all three sections agree on how many
    samples this report covers. `sample_photos` (already-cropped JPEG
    bytes, keyed by the same label this function writes into column 1)
    is optional - unmatched or missing samples simply keep blank photo
    cells, same as before this existed. The After-Test caption is never
    taken from `sample_photos` - it's always this same sample's own
    `remarks` (identical text to what Section 2 shows), matching how
    real reports caption that column with the leak/pass outcome."""
    if len(doc.tables) < 8:
        return
    results = _effective_sample_results(data) or []
    part_no = data.seal_part_no
    table = doc.tables[7]
    _resize_table(table, len(results) + 1)
    sample_photos = sample_photos or {}
    for idx, result in enumerate(results, start=1):
        label = _sample_photo_label(result.sample_no, part_no)
        _set_row_values(table.rows[idx], [label, "", ""])
        entry = sample_photos.get(label) or SamplePhotoEntry()
        row_cells = table.rows[idx].cells
        _set_cell_picture(row_cells[1], entry.before_photo, entry.before_caption)
        _set_cell_picture(row_cells[2], entry.after_photo, result.remarks)


def _sample_photo_label(sample: str, part_no: str) -> str:
    """"70521 Sample #01" style label for the Section 5 photos table -
    spelled out, unlike the "S#01" suffix shorthand used everywhere else
    in the report, per how this table is meant to read. Mirrors
    `_sample_column_label`'s part-number lookup: a sample string that
    already carries its own part number (multi-part-number reports, e.g.
    a combined 71264 & 71174 cold test) wins over the report's single
    `data.seal_part_no`, so each sample keeps its own part number
    instead of collapsing onto whichever one was extracted first."""
    suffix = _sample_suffix(sample)
    match = re.search(r"(\d+)$", suffix)
    number = match.group(1) if match else "01"
    sample_part = _part_no_from_sample(sample) or part_no
    return f"{sample_part} Sample #{number}".strip()


def _set_cell_picture(cell, photo_jpeg: bytes | None, caption: str = "") -> None:
    if not photo_jpeg:
        return
    width = Inches(2.6)
    if cell.width and cell.width < width + Inches(0.15):
        width = max(Inches(0.5), cell.width - Inches(0.15))
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    picture = run.add_picture(io.BytesIO(photo_jpeg), width=width)
    _apply_picture_frame(picture, SAMPLE_PHOTO_BORDER_COLOR)

    caption = (caption or "").strip()
    if caption:
        # No explicit font/size override: table body text throughout this
        # template (e.g. the Sample No. table, this same photos table's
        # own label column) never overrides the run font either, so
        # leaving this unset inherits the same default the rest of the
        # table already uses instead of picking an arbitrary size.
        caption_paragraph = cell.add_paragraph()
        caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_paragraph.add_run(caption)


# Colors observed in the report format: the template's own section 3.2
# example photo uses a blue frame; the real edited reports' Section 5
# Before/After Test photos consistently use a black one (71595, 71529,
# 71710 all match; only one of four sampled reports used blue there).
SEAL_PHOTO_BORDER_COLOR = "0F58D6"
SAMPLE_PHOTO_BORDER_COLOR = "000000"
_PICTURE_FRAME_BORDER_WIDTH_EMU = 9525  # 0.75pt, matches every sampled photo


def _apply_picture_frame(inline_shape, border_color: str) -> None:
    """Gives a just-inserted picture the same rounded-rectangle-with-
    colored-border frame every photo in the reference report format
    uses (section 3.2's identification photo, Section 5's Before/After
    Test photos) - `add_picture` alone produces a plain, unstyled
    rectangle that doesn't match."""
    pic_elements = inline_shape._inline.xpath(".//pic:pic")
    if not pic_elements:
        return
    spPr = pic_elements[0].xpath("./pic:spPr")
    if not spPr:
        return
    spPr = spPr[0]

    geometry = spPr.find(qn("a:prstGeom"))
    if geometry is not None:
        geometry.set("prst", "roundRect")
        if geometry.find(qn("a:avLst")) is None:
            geometry.append(OxmlElement("a:avLst"))

    line = OxmlElement("a:ln")
    line.set("w", str(_PICTURE_FRAME_BORDER_WIDTH_EMU))
    line.set("cap", "flat")
    line.set("cmpd", "sng")
    line.set("algn", "ctr")
    solid_fill = OxmlElement("a:solidFill")
    color = OxmlElement("a:srgbClr")
    color.set("val", border_color)
    solid_fill.append(color)
    line.append(solid_fill)
    dash = OxmlElement("a:prstDash")
    dash.set("val", "solid")
    line.append(dash)
    spPr.append(line)


def detect_sample_labels(
    project_spec_path: Path,
    inspection_sheet_path: Path,
    monitoring_sheet_paths: list[Path] | None = None,
) -> list[str]:
    """Parses just the Project Specification + Test Inspection sheet (and
    monitoring sheets, if already attached) far enough to know the sample
    count and labels this report will use under Sections 2/4.3/5 - lets a
    caller render the right number of Before/After photo slots before the
    user commits to generating the full report. Raises whatever
    `_extract_data` raises on unreadable input; callers should catch and
    surface that as a normal validation error, not a crash."""
    data = _extract_data(project_spec_path, inspection_sheet_path, monitoring_sheet_paths or [])
    samples = [result.sample_no for result in _effective_sample_results(data)] or ["S#01"]
    return [_sample_photo_label(sample, data.seal_part_no) for sample in samples]


def _set_test_description_table(doc: Document, data: ExtractedFinalReportData) -> None:
    if len(doc.tables) < 6:
        return

    samples = _final_sample_numbers(data)
    table = doc.tables[5]
    # The template's sample columns carry leftover horizontal merges from
    # manual width tweaking in Word (e.g. gridSpan="2"/"5" on various
    # rows). python-docx's row.cells returns the *same* cell object for
    # every position a merge spans, so writing distinct per-sample values
    # by index would silently collide into whichever write happens last.
    # Unmerge first so every column index maps to an independent cell.
    for row in table.rows:
        _unmerge_row_cells(row)
    _resize_table_columns(table, 2 + len(samples))

    sample_headers = [_sample_column_label(sample, data.seal_part_no) for sample in samples]
    row_values = [
        ["Sample Designation.", "Specification as per OSTR/Drawing", *sample_headers],
        ["Test Request Form #", data.test_request_form_no, *([data.test_request_form_no] * len(samples))],
        ["Total Duration [Hrs.]", _hours_text(data.target_hours), *[_hours_text(result.actual_hours or data.actual_hours) for result in _results_for_samples(data, samples)]],
        ["Speed [rpm]", "Refer Test cycle (after this table)", *["Refer Test Cycle" for _ in samples]],
        ["Direction of Rotation", data.rotation_direction, *[data.rotation_direction for _ in samples]],
        _spec_row("Dynamic Run out [mm]", data.dro, samples),
        _spec_row("Static Run out [mm]", data.stbm, samples),
        _spec_row("Seal Cock [mm]", data.seal_cock, samples),
    ]
    # Reciprocation is only measured/reported when the requestor's test
    # specification actually asks for it.
    if data.reciprocation.spec:
        row_values.append(_spec_row("Reciprocation [mm]", data.reciprocation, samples))
    row_values.extend([
        _spec_row("Shaft Surface Finish (µm Ra)", data.shaft_surface_finish, samples),
        _spec_row("Shaft Diameter (mm)", data.shaft_diameter, samples),
        _spec_row("Shaft hardness (HRC)", data.shaft_hardness, samples),
        _spec_row("Housing Diameter (mm)", data.housing_diameter, samples),
        _spec_row("Housing Surface Finish (µm Ra)", data.housing_surface_finish, samples),
        _spec_row("Fluid", data.fluid, samples),
        ["Temperature (°C)", "Refer Test cycle (after this table)", *["Refer Test cycle (after this table)" for _ in samples]],
        _spec_row("Fluid level", data.fluid_level, samples),
        [
            "Oil change interval",
            _hours_text(data.oil_change_interval.spec),
            *[_hours_text_short(data.oil_change_interval.value_for(sample, data.oil_change_interval.spec)) for sample in samples],
        ],
    ])

    for row_idx, values in enumerate(row_values):
        if row_idx >= len(table.rows):
            table.add_row()
        _set_row_values(table.rows[row_idx], values)


def _spec_row(label: str, value: TestSpecValue, samples: list[str]) -> list[str]:
    """Build one Test description table row: the spec column always shows
    the Project Specification's value; each sample column shows that
    sample's actual measured value where one was found, else the same
    spec value (used for constants like Fluid / Fluid level)."""
    return [label, value.spec, *[value.value_for(sample, value.spec) for sample in samples]]


def _set_test_dates_table(doc: Document, data: ExtractedFinalReportData) -> None:
    if len(doc.tables) < 7:
        return
    entries = data.test_dates or []
    if not entries:
        return

    table = doc.tables[6]
    # Position/Machine Identification only apply to the multi-position
    # GB1GTMC011 machine; leave the table at its original 3 columns when
    # none of the samples have that Section C position code at all.
    has_position = any(entry.position for entry in entries)
    _resize_table_columns(table, 5 if has_position else 3)
    if has_position:
        header = table.rows[0].cells
        if len(header) >= 5:
            if not header[3].text.strip():
                header[3].text = "Position"
            if not header[4].text.strip():
                header[4].text = "Machine Identification"

    _resize_table(table, len(entries) + 1)
    for idx, entry in enumerate(entries, start=1):
        values = [
            _sample_column_label(entry.sample_no, data.seal_part_no),
            entry.start_date,
            entry.end_date,
        ]
        if has_position:
            values.extend([entry.position, entry.machine_identification])
        _set_row_values(table.rows[idx], values)


def _set_pre_post_measurement_tables(doc: Document, data: ExtractedFinalReportData) -> None:
    """Fill the "Pre and Post Test measurement results" tables (Section 5)
    from Section B of the Inspection & Execution sheet, rather than leaving
    the template's own worked example (Main lip Dia / Metal Outer Dia /
    ... for a 3-sample project) in every report. Which description rows
    appear varies project to project, so only the rows the technician
    actually filled in are used - see _extract_section_b_rows."""
    if len(doc.tables) < 10:
        return
    samples = _final_sample_numbers(data)
    if not samples:
        return
    measurements = data.section_b_measurements or {}
    if not measurements:
        return
    specs = data.section_b_specs or {}

    sample_suffixes = [_sample_suffix(sample) for sample in samples]
    _fill_measurement_table(doc.tables[8], sample_suffixes, measurements, specs, "pre", "Pre-Test", data.seal_part_no, data.section_b_remarks)
    _fill_measurement_table(doc.tables[9], sample_suffixes, measurements, specs, "post", "Post-Test", data.seal_part_no, data.section_b_remarks)


def _fill_measurement_table(
    table,
    sample_suffixes: list[str],
    measurements: dict[str, dict[str, tuple[str, str]]],
    specs: dict[str, str],
    which: str,
    banner_label: str,
    part_no: str,
    remarks_text: str,
) -> None:
    labels = list(measurements.keys())
    column_count = 4 + len(sample_suffixes)
    row_count = 2 + len(labels) + (1 if remarks_text else 0)

    # The template's sample columns carry a horizontal merge on the banner
    # row (one "Pre-Test Measurement - <part>" cell spanning every sample
    # column). Unmerge before resizing so each column index maps to an
    # independent cell, same reasoning as _set_test_description_table.
    for row in table.rows:
        _unmerge_row_cells(row)
    _resize_table_columns(table, column_count)
    _resize_table(table, row_count)

    fixed_labels = ["Sl. No", "Description", "Spec \n(as per Drwg. / OSTR)", "Measuring Equipment"]
    banner_cells = table.rows[0].cells
    for idx, value in enumerate(fixed_labels):
        if idx < len(banner_cells):
            banner_cells[idx].text = value
    banner_text = f"{banner_label} Measurement - {part_no}".strip(" -")
    for idx in range(len(fixed_labels), len(banner_cells)):
        banner_cells[idx].text = banner_text

    _set_row_values(table.rows[1], [*fixed_labels, *sample_suffixes])

    row_idx = 2
    for sl_no, label in enumerate(labels, start=1):
        by_sample = measurements[label]
        values = [str(sl_no), label, specs.get(label) or "NA", "NA"]
        for suffix in sample_suffixes:
            pair = by_sample.get(suffix)
            values.append((pair[0] if which == "pre" else pair[1]) if pair else "")
        _set_row_values(table.rows[row_idx], values)
        row_idx += 1

    if remarks_text:
        row = table.rows[row_idx]
        for cell in row.cells:
            cell.text = ""
        if len(row.cells) > 1:
            row.cells[0].text = "Remarks"
            row.cells[1].text = remarks_text
        elif row.cells:
            row.cells[0].text = f"Remarks: {remarks_text}"


def _set_footer_project_no(doc: Document, project_no: str) -> None:
    for section in doc.sections:
        for footer in (section.footer, section.first_page_footer, section.even_page_footer):
            if not footer.paragraphs:
                footer.add_paragraph()
            touched = False
            for p in footer.paragraphs:
                if "Request" in p.text or PROJECT_NO_RE.search(p.text):
                    _set_paragraph_text(p, f"Request No. {project_no}")
                    touched = True
            if not touched:
                footer.paragraphs[0].text = f"Request No. {project_no}"


def _refresh_toc_page_numbers(doc: Document) -> None:
    heading_pages = {
        "Objective": "5",
        "Conclusions": "5",
        "Test samples identification": "5",
        "Date of receipt of the samples": "5",
        "Identification of the samples": "5",
        "Description and condition of the test samples": "6",
        "Test description": "6",
        "Description/specification of the test": "6",
        "Test and measuring equipment.": "9",
        "Date the test was performed.": "9",
        "Test results": "9",
    }
    for p in doc.paragraphs:
        text = p.text.strip()
        for heading, page in heading_pages.items():
            if text.endswith(f"\t{heading}\t5") or text.endswith(f"\t{heading}\t6") or text.endswith(f"\t{heading}\t9"):
                parts = text.split("\t")
                parts[-1] = page
                _set_paragraph_text(p, "\t".join(parts))


def _patch_textbox_title(path: Path, title: str, author: str, project_no: str, report_date: str, approver: str) -> None:
    with tempfile.TemporaryDirectory() as td:
        temp_dir = Path(td)
        with zipfile.ZipFile(path) as zin:
            zin.extractall(temp_dir)

        document_xml = temp_dir / "word/document.xml"
        _patch_document_xml(document_xml, title, author, project_no, report_date, approver)

        for xml_path in (temp_dir / "word").glob("*.xml"):
            if xml_path.name.startswith(("header", "footer")):
                text = xml_path.read_text(encoding="utf-8")
                text = re.sub(r"TR25-0009-BTS", project_no, text)
                text = re.sub(r"TR2</w:t>.*?-BTS</w:t>", project_no, text, flags=re.DOTALL)
                xml_path.write_text(text, encoding="utf-8")

        with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in temp_dir.rglob("*"):
                if item.is_file():
                    zout.write(item, item.relative_to(temp_dir).as_posix())


def _patch_document_xml(xml_path: Path, title: str, author: str, project_no: str, report_date: str, approver: str) -> None:
    from lxml import etree

    root = etree.fromstring(xml_path.read_bytes())
    text_nodes = root.xpath(".//*[local-name()='t']")
    nonempty = [node for node in text_nodes if (node.text or "").strip()]
    if len(nonempty) > 25:
        # Cover page textbox text in the reference template is split into many
        # runs. Put the new values in the first run of each visual line and
        # clear the remaining old fragments so the textbox positions survive.
        for start, end, value in (
            (0, 5, title),
            (5, 11, title),
            (11, 12, author),
            (16, 21, project_no),
            (21, 22, report_date),
            (22, 23, title),
        ):
            nonempty[start].text = value
            for node in nonempty[start + 1:end]:
                node.text = ""
    for node in text_nodes:
        text = node.text or ""
        text = text.replace("TR25-0009-BTS", project_no)
        text = text.replace("01 April 2025", report_date)
        text = text.replace("Anoop N", author)
        text = text.replace("Dhayalan K", approver)
        node.text = text
    xml_path.write_bytes(etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True))


def _set_paragraph_text(paragraph, text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    for run in paragraph.runs:
        run.text = ""
    paragraph.runs[0].text = text


def _ensure_rows(table, count: int) -> None:
    """`table.add_row()` builds a bare new `<w:tr>` - it does not copy the
    per-cell `tcPr` (borders, shading, ...) that this template relies on,
    since these tables define borders per-cell rather than once at the
    table level (`w:tblBorders`). Left alone, every row this generator
    adds beyond the template's original row count would render with no
    visible cell borders at all. Copying the last existing row's cell
    formatting onto each new row keeps them visually identical to the
    template, matching the reference report format."""
    template_tr = table.rows[-1]._tr if table.rows else None
    while len(table.rows) < count:
        new_row = table.add_row()
        if template_tr is not None:
            _copy_row_cell_formatting(new_row, template_tr)


def _copy_row_cell_formatting(new_row, template_tr) -> None:
    template_cells = template_tr.findall(qn("w:tc"))
    new_cells = new_row._tr.findall(qn("w:tc"))
    for new_tc, template_tc in zip(new_cells, template_cells):
        template_tc_pr = template_tc.find(qn("w:tcPr"))
        if template_tc_pr is None:
            continue
        existing_tc_pr = new_tc.find(qn("w:tcPr"))
        if existing_tc_pr is not None:
            new_tc.remove(existing_tc_pr)
        new_tc.insert(0, copy.deepcopy(template_tc_pr))


def _resize_table(table, count: int) -> None:
    _ensure_rows(table, count)
    while len(table.rows) > count:
        table._tbl.remove(table.rows[-1]._tr)


def _unmerge_row_cells(row) -> None:
    """Split any horizontally-merged cell (w:gridSpan > 1) into that many
    independent cells, so row.cells afterward maps 1:1 to physical grid
    columns instead of python-docx returning the same cell object for
    every position the merge spans (which would silently collapse
    per-column writes to whichever one happens last)."""
    for tc in list(row._tr.findall(qn("w:tc"))):
        tc_pr = tc.find(qn("w:tcPr"))
        grid_span = tc_pr.find(qn("w:gridSpan")) if tc_pr is not None else None
        span = int(grid_span.get(qn("w:val"))) if grid_span is not None else 1
        if span <= 1:
            continue
        tc_pr.remove(grid_span)
        insert_after = tc
        for _ in range(span - 1):
            new_tc = OxmlElement("w:tc")
            new_tc.append(OxmlElement("w:tcPr"))
            new_tc.append(OxmlElement("w:p"))
            insert_after.addnext(new_tc)
            insert_after = new_tc


def _resize_table_columns(table, count: int) -> None:
    if count <= 0 or not table.rows:
        return
    for row in table.rows:
        while len(row.cells) > count:
            row._tr.remove(row.cells[-1]._tc)
        while len(row.cells) < count:
            tc = OxmlElement("w:tc")
            tc_pr = OxmlElement("w:tcPr")
            tc.append(tc_pr)
            tc.append(OxmlElement("w:p"))
            row._tr.append(tc)

    grid = table._tbl.tblGrid
    if grid is not None:
        while len(grid.gridCol_lst) > count:
            grid.remove(grid.gridCol_lst[-1])
        while len(grid.gridCol_lst) < count:
            grid_col = OxmlElement("w:gridCol")
            grid_col.set(qn("w:w"), "1800")
            grid.append(grid_col)


def _add_person(people: list[tuple[str, str, str]], name: str, role: str, location: str) -> None:
    clean_name = (name or "").strip()
    if not clean_name:
        return
    if any(existing[0].casefold() == clean_name.casefold() for existing in people):
        return
    people.append((clean_name, role, location))


def _set_row_values(row, values: list[str]) -> None:
    cells = row.cells
    if not cells:
        return
    if len(cells) >= len(values):
        for idx, value in enumerate(values):
            cells[idx].text = value
        return
    cells[0].text = values[0]
    if len(cells) > 1:
        cells[1].text = " | ".join(value for value in values[1:] if value)


def _set_cell_dropdown(cell, value: str, options: list[str]) -> None:
    selected = value if value in options else options[2]
    tc = cell._tc
    for child in list(tc):
        if child.tag != qn("w:tcPr"):
            tc.remove(child)

    sdt = OxmlElement("w:sdt")
    sdt_pr = OxmlElement("w:sdtPr")

    alias = OxmlElement("w:alias")
    alias.set(qn("w:val"), "Result")
    sdt_pr.append(alias)

    tag = OxmlElement("w:tag")
    tag.set(qn("w:val"), "Result")
    sdt_pr.append(tag)

    dropdown = OxmlElement("w:dropDownList")
    for option in options:
        item = OxmlElement("w:listItem")
        item.set(qn("w:displayText"), option)
        item.set(qn("w:value"), option)
        dropdown.append(item)
    sdt_pr.append(dropdown)

    sdt_content = OxmlElement("w:sdtContent")
    paragraph = OxmlElement("w:p")
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = selected
    run.append(text)
    paragraph.append(run)
    sdt_content.append(paragraph)

    sdt.append(sdt_pr)
    sdt.append(sdt_content)
    tc.append(sdt)


def _first_match(pattern: re.Pattern[str], text: str) -> str:
    match = pattern.search(text or "")
    return match.group(0) if match else ""


def _extract_label_value(text: str, labels: list[str]) -> str:
    for label in labels:
        pattern = re.compile(rf"{re.escape(label)}\s*[:|\-]?\s*([^\n|]+)", re.IGNORECASE)
        match = pattern.search(text or "")
        if match:
            value = match.group(1).strip()
            if value and value.lower() != label.lower():
                return value
    return ""


def _extract_requester_name(text: str) -> str:
    """The requester's name and the adjacent "Project Leader" field are
    often on the same extracted line ("Requested By Manjunath H R Project
    Leader Anoop N") with no delimiter between them, so a plain
    capture-to-end-of-line would swallow both; stop right before
    "Project Leader" instead."""
    match = re.search(
        r"Requested\s+By\s*[:|\-]?\s*([^\n|]+?)(?=\s+Project\s+Leader\b|\n|$)",
        text or "",
        re.IGNORECASE,
    )
    if match:
        value = match.group(1).strip()
        if value:
            return value
    return _extract_label_value(text, ["Requester", "Requestor"])


def _extract_project_title(text: str, project_no: str) -> str:
    lines = [line.strip(" \t|") for line in (text or "").splitlines() if line.strip()]
    for line in lines:
        if project_no and project_no in line:
            continue
        lowered = line.lower()
        if "project specification" in lowered or "final report" in lowered:
            continue
        if any(word in lowered for word in ("test", "seal", "endurance", "dust", "slurry")) and len(line) > 20:
            return line
    return ""


_NEXT_SECTION_HEADINGS = {
    "scope", "test conditions", "test procedure", "test cycle", "requirements",
    "acceptance criteria", "acceptance criterion", "sample details", "test setup",
    "references", "test request form", "test request form #", "sample identification",
    "definitions", "abbreviations", "procedure for monitoring", "disclaimer",
    "product drawing", "decision rule",
}


def _extract_section_line(text: str, heading: str) -> str:
    """Return the full section body verbatim, spanning multiple wrapped
    lines when the source (e.g. a PDF) splits one logical paragraph into
    several text lines. Stops at a blank-line gap or the next known
    section heading so the objective is copied exactly, not truncated."""
    lines = [line.strip() for line in (text or "").splitlines()]
    for idx, line in enumerate(lines):
        if line.strip(":").casefold() == heading.casefold():
            collected: list[str] = []
            for candidate in lines[idx + 1:idx + 25]:
                clean = candidate.strip(":")
                if not clean:
                    if collected:
                        break
                    continue
                if clean.casefold() in _NEXT_SECTION_HEADINGS or clean.casefold() == heading.casefold():
                    break
                collected.append(candidate)
            if collected:
                return " ".join(collected)
    return ""


_BULLET_MARKER_RE = re.compile(r"^[➢•▪◦‣∙]\s*|^[-*]\s+")


def _extract_bullet_lines(text: str, heading: str, max_lines: int = 150) -> list[str]:
    """Return each bullet under `heading` as a separate list entry
    (unlike _extract_section_line, which joins everything into one
    string) so multi-bullet sections like Acceptance Criteria, Procedure
    for monitoring, or Test Procedure can be copied bullet-for-bullet.

    When the source marks bullets with an explicit symbol (➢, •, ...),
    that marker is the authoritative boundary and is stripped from the
    result; any other line encountered while not actively continuing a
    bullet is treated as noise and skipped - this is what lets Test
    Procedure's bullets survive being interspersed with an embedded
    measurement table's flattened, garbled text. Without any markers in
    the section, boundaries fall back to the source line's leading
    indent or a completed sentence followed by a new capitalised line."""
    raw_lines = (text or "").splitlines()
    start_idx = None
    for idx, line in enumerate(raw_lines):
        if line.strip().strip(":").casefold() == heading.strip(":").casefold():
            start_idx = idx + 1
            break
    if start_idx is None:
        return []

    window = raw_lines[start_idx:start_idx + max_lines]
    uses_markers = any(_BULLET_MARKER_RE.match(line.strip()) for line in window)

    bullets: list[str] = []
    current = ""
    accumulating = False
    for raw_line in window:
        clean = raw_line.strip()
        if not clean:
            continue
        if clean.strip(":").casefold() in _NEXT_SECTION_HEADINGS:
            break

        marker_match = _BULLET_MARKER_RE.match(clean)
        if marker_match:
            if current:
                bullets.append(current.strip())
            current = clean[marker_match.end():].strip()
            accumulating = True
            continue

        if uses_markers:
            complete = current.rstrip().endswith((".", ":", ";"))
            if accumulating and current and not complete and re.match(r"^[a-z0-9(]", clean):
                current = f"{current} {clean}".strip()
            # Anything else here (an embedded table's flattened rows, a
            # stray heading fragment, ...) is noise between bullets.
            continue

        indented = raw_line[:1] in (" ", "\t")
        starts_new_sentence = bool(re.match(r"^[A-Z0-9]", clean)) and current.rstrip().endswith((".", ":", ";"))
        if current and (indented or starts_new_sentence):
            bullets.append(current.strip())
            current = clean
        else:
            current = f"{current} {clean}".strip()
    if current:
        bullets.append(current.strip())
    return bullets


def _extract_objective_sentence(text: str) -> str:
    match = re.search(r"(To conduct[^\n]+)", text or "", flags=re.IGNORECASE)
    return match.group(1).strip() if match else ""


def _extract_target_hours(text: str) -> str:
    patterns = [
        r"Total\s+test\s+duration[^\d]*(\d+)\s*hours",
        r"Target\s+Hours?[^\d]*(\d+)",
        r"Total\s+Duration\s*\[?Hrs?\.?\]?[^\d]*(\d+)",
        # Some specs phrase it as "duration is 40 cycles (960 hours)" - the
        # number immediately before "hours" isn't the one right after
        # "duration" in that case, so look anywhere on the same line.
        r"Total\s+test\s+duration[^\n]*?(\d+)\s*hours",
    ]
    for pattern in patterns:
        match = re.search(pattern, text or "", re.IGNORECASE)
        if match:
            return match.group(1)
    return ""


def _extract_actual_hours(text: str, target_hours: str = "") -> str:
    match = re.search(
        r"(?:Completed|Actual|Run|Total|Hours\s+completed)\s+(?:Test\s+)?(?:Hours?|Duration|completed)?[^\d]*(\d+)",
        text or "",
        re.IGNORECASE,
    )
    if match:
        return match.group(1)

    # OCR of a handwritten daily monitoring sheet loses the "Completed: N"
    # phrasing (each day's row is just a bare number under a "Hours
    # completed" column header). The log accumulates toward the planned
    # target, so the closest-to-target plausible number is the best
    # available reading of the final completed hours.
    target = _first_number(target_hours)
    if target is None:
        return ""
    candidates = [int(value) for value in re.findall(r"\b\d{2,4}\b", text or "")]
    plausible = [value for value in candidates if 0 < value <= target + 20]
    return str(max(plausible)) if plausible else ""


def _extract_test_request_form_no(objective_text: str, full_text: str, fallback: str) -> str:
    patterns = [
        r"\b(Test\s+Request\s+form\s*#?\s*[:|\-]?\s*[A-Z0-9./_-]+)\b",
        r"\b(Request\s+(?:No|Number)\.?\s*[:|\-]?\s*[A-Z0-9./_-]+)\b",
        r"\b(TR\d{2}-\d{4}-BTS)\b",
    ]
    # The request number is defined in the Objective section of the project
    # specification; prefer it there before falling back to a whole-document
    # search so an unrelated number elsewhere in the spec isn't picked up.
    for source in (objective_text, full_text):
        for pattern in patterns:
            match = re.search(pattern, source or "", re.IGNORECASE)
            if match:
                return " ".join(match.group(1).split())
    return fallback


def _extract_rotation_direction(text: str) -> str:
    cycle_text = _text_after_heading(text, "Test cycle", max_lines=35) or text
    has_cw = re.search(r"\bCW\b|clockwise", cycle_text or "", re.IGNORECASE) is not None
    has_ccw = re.search(r"\bCCW\b|counter\s*clockwise|anti\s*clockwise", cycle_text or "", re.IGNORECASE) is not None
    if has_cw and has_ccw:
        return "Bi-directional"
    if has_cw:
        return "CW"
    if has_ccw:
        return "CCW"
    return "Refer Test Cycle"


def _extract_sample_receipt_date(text: str) -> str:
    label_value = _extract_label_value(text, ["Sample Receipt date", "Sample Receipt Date", "Receipt date", "Received date"])
    date = _first_date(label_value)
    if date:
        return date
    for line in (text or "").splitlines():
        if "receipt" in line.lower() or "received" in line.lower():
            date = _first_date(line)
            if date:
                return date
    return ""


def _extract_seal_part_no(*texts: str) -> str:
    combined = "\n".join(text for text in texts if text)
    patterns = [
        r"\b(?:Seal\s*)?(?:Part\s*)?No\.?\s*[:|\-]?\s*(\d{5,8})\b",
        r"\b(\d{5,8})\s*[–-]\s*(?:ID|S#|Sample|Seal)(?=\b|\d)",
        # Titles like "70521 seal Endurance Test..." or "71261Seal ..." name
        # the part number directly before the word "seal", with no ID/S#/
        # Sample/dash marker at all.
        r"\b(\d{5,8})\s*seals?\b",
    ]
    for pattern in patterns:
        match = re.search(pattern, combined, re.IGNORECASE)
        if match:
            return match.group(1)
    return ""


def _extract_sample_numbers(text: str) -> list[str]:
    matches = re.findall(r"(?:\d{5,8}\s*[–-]\s*)?S#\s*0?\d+", text or "", flags=re.IGNORECASE)
    cleaned: list[str] = []
    for match in matches:
        value = _normalize_sample_no(match)
        if value not in cleaned:
            cleaned.append(value)
    return cleaned[:12]


def _extract_remarks(text: str) -> list[str]:
    remarks: list[str] = []
    for line in (text or "").splitlines():
        lowered = line.lower()
        if "sl. no" in lowered or "sample #" in lowered or "target hours" in lowered:
            continue
        if re.search(r"\b(pass|fail)\b", line, re.IGNORECASE) or "remark" in lowered:
            clean = " ".join(_remark_text_from_row(line).split())
            if clean and clean not in remarks:
                remarks.append(clean)
    return remarks[:12]


def _remark_text_from_row(line: str) -> str:
    """Pull just the remark wording out of a source line.

    `_extract_docx_text` flattens each table row into a single
    "cell | cell | cell" line, so a results-table row arrives here as
    e.g. "1 | 71261 - S#01 | 960 | 960 | PASS | Zero Leakage observed".
    Keeping the whole row would put that entire pipe-joined string into
    the report's Remarks column (and, since Section 5 captions its
    After-Test photos with the same field, into the photo caption too).
    Only the descriptive cell is wanted, so drop the columns that are
    pure bookkeeping - serial numbers, sample designations, hour counts,
    and the bare PASS/FAIL verdict - and keep the last one with actual
    wording left in it. Lines without any "|" are returned unchanged."""
    if "|" not in line:
        return line
    for cell in reversed([part.strip() for part in line.split("|")]):
        if not cell:
            continue
        if re.fullmatch(r"(?i)(pass|fail)", cell):
            continue
        # Bare numbers/hours ("960", "960 Hrs.") and sample or part
        # designations ("1", "71261 - S#01", "S#02") carry no wording.
        if re.fullmatch(r"[\d.\s]+(?:hrs?\.?|hours?)?", cell, re.IGNORECASE):
            continue
        if re.fullmatch(r"(?i)[\d\s\-–]*s\s*#\s*\d+", cell):
            continue
        if not re.search(r"[A-Za-z]{3}", cell):
            continue
        return cell
    return ""


def _extract_suspension_note(text: str) -> str:
    """A verbatim reason for why fewer samples were tested than requested
    (e.g. "Test got suspended after the 1st sample test, as per the
    requester request"), when one is present in the source documents.
    Absence of this note is not itself meaningful - callers fall back to a
    generic phrase when it comes back empty."""
    for line in (text or "").splitlines():
        if "suspend" in line.lower():
            note = re.sub(r"^\s*note\s*[:.\-]\s*", "", line.strip(), flags=re.IGNORECASE)
            note = " ".join(note.split())
            if note and note[-1] not in ".!?":
                note += "."
            return note[0].upper() + note[1:] if note else ""
    return ""


def _extract_result(remark: str) -> str:
    if re.search(r"\bfail\b", remark or "", re.IGNORECASE):
        return "FAIL"
    if re.search(r"\bpass\b", remark or "", re.IGNORECASE):
        return "PASS"
    return ""


def _build_sample_results(
    data: ExtractedFinalReportData,
    inspection_text: str,
    monitoring_texts: list[str],
) -> list[SampleResult]:
    samples = data.sample_numbers or _sample_numbers_from_count(_extract_sample_count_from_text(data.objective))
    if not samples:
        samples = ["S#01"]

    remarks_by_sample = _extract_remarks_by_sample(inspection_text)
    monitoring_by_sample = _match_monitoring_texts_to_samples(samples, monitoring_texts)
    results: list[SampleResult] = []

    for idx, sample in enumerate(samples):
        monitor_text = monitoring_by_sample.get(sample, "")
        remark = remarks_by_sample.get(sample)
        if remark is None and data.remarks and idx < len(data.remarks):
            remark = data.remarks[idx]
        remark = remark or _extract_monitoring_remark(monitor_text)
        actual_hours = _extract_actual_hours(monitor_text, data.target_hours) or data.actual_hours
        result = _derive_result(data.target_hours, actual_hours, remark)
        results.append(
            SampleResult(
                sample_no=sample,
                target_hours=data.target_hours,
                actual_hours=actual_hours,
                result=result,
                remarks=_clean_result_from_remark(remark),
            )
        )
    return results


def _extract_remarks_by_sample(text: str) -> dict[str, str]:
    remarks: dict[str, str] = {}
    for line in (text or "").splitlines():
        sample_match = re.search(r"(?:\d{5,8}\s*[–-]\s*)?S#\s*0?\d+", line, re.IGNORECASE)
        if not sample_match:
            continue
        lowered = line.lower()
        if "remark" not in lowered and "leak" not in lowered and not re.search(r"\b(pass|fail)\b", line, re.IGNORECASE):
            continue
        sample = _normalize_sample_no(sample_match.group(0))
        clean = " ".join(_remark_text_from_row(line).split())
        if clean:
            remarks[sample] = clean
    return remarks


def _match_monitoring_texts_to_samples(samples: list[str], monitoring_texts: list[str]) -> dict[str, str]:
    matched: dict[str, str] = {}
    normalized_samples = {_sample_suffix(sample): sample for sample in samples}
    for idx, text in enumerate(monitoring_texts):
        sample = ""
        for suffix, full_sample in normalized_samples.items():
            if suffix and re.search(rf"S#\s*0?{re.escape(suffix[-2:])}\b", text or "", re.IGNORECASE):
                sample = full_sample
                break
        if not sample and idx < len(samples):
            sample = samples[idx]
        if sample:
            matched[sample] = text
    return matched


def _extract_monitoring_remark(text: str) -> str:
    for line in (text or "").splitlines():
        lowered = line.lower()
        if "leak" in lowered or "remark" in lowered or re.search(r"\b(pass|fail)\b", line, re.IGNORECASE):
            clean = " ".join(_remark_text_from_row(line).split())
            if clean:
                return clean
    return ""


def _derive_result(target_hours: str, actual_hours: str, remark: str) -> str:
    explicit = _extract_result(remark)
    if explicit:
        return explicit
    target = _first_number(target_hours)
    actual = _first_number(actual_hours)
    if target is not None and actual is not None and target == actual:
        return "PASS"
    return "No results defined"


def _clean_result_from_remark(remark: str) -> str:
    return re.sub(r"\b(pass|fail)\b", "", remark or "", flags=re.IGNORECASE).strip(" -:;")


def _first_number(text: str) -> int | None:
    match = re.search(r"\d+", text or "")
    return int(match.group(0)) if match else None


def _first_date(text: str) -> str:
    match = re.search(
        r"\b(\d{1,2}[./-]\d{1,2}[./-]\d{2,4}|\d{1,2}(?:st|nd|rd|th)?\s+[A-Za-z]+\s+\d{4})\b",
        text or "",
        re.IGNORECASE,
    )
    return match.group(1) if match else ""


def _normalize_sample_no(value: str) -> str:
    part_match = re.search(r"\b\d{5,8}\b", value or "")
    sample_match = re.search(r"S#\s*0?(\d+)", value or "", re.IGNORECASE)
    if not sample_match:
        return value.strip()
    suffix = f"S#{int(sample_match.group(1)):02d}"
    return f"{part_match.group(0)} - {suffix}" if part_match else suffix


def _sample_suffix(value: str) -> str:
    match = re.search(r"S#\s*0?(\d+)", value or "", re.IGNORECASE)
    return f"S#{int(match.group(1)):02d}" if match else value.strip()


def _extract_sample_count_from_text(text: str) -> int:
    match = re.search(r"\bon\s+(\d+)\s+samples?\b|\b(\d+)\s+samples?\b", text or "", re.IGNORECASE)
    if not match:
        return 0
    value = match.group(1) or match.group(2)
    return int(value)


def _extract_tested_sample_count(text: str) -> int:
    section_c = _text_after_heading(text, "Section C", max_lines=80) or _text_after_heading(text, "C.", max_lines=80)
    search_text = section_c or text
    patterns = [
        r"\b(?:actual\s+)?(?:no\.?\s+of\s+)?samples?\s+(?:tested|inspected|received)\s*[:|\-]?\s*(\d+)\b",
        r"\b(\d+)\s+samples?\s+(?:tested|inspected|received)\b",
        r"\bno\.?\s+of\s+samples?\s*[:|\-]?\s*(\d+)\b",
        r"\bsamples?\s+quantity\s*[:|\-]?\s*(\d+)\b",
        r"\bquantity\s+of\s+samples?\s*[:|\-]?\s*(\d+)\b",
    ]
    for pattern in patterns:
        match = re.search(pattern, search_text or "", re.IGNORECASE)
        if match:
            return int(match.group(1))
    sample_indexes = re.findall(r"\bS#\s*0?(\d+)\b", search_text or "", re.IGNORECASE)
    if sample_indexes:
        return max(int(value) for value in sample_indexes)
    return 0


def _select_tested_sample_numbers(text: str, part_no: str, requested_count: int, tested_count: int) -> list[str]:
    samples = _extract_sample_numbers(text)
    count = tested_count or requested_count or len(samples)
    if count:
        return [_sample_no_for_index(part_no, idx) for idx in range(1, count + 1)]
    return samples


def _sample_numbers_from_count(count: int) -> list[str]:
    if count <= 0:
        return []
    return [f"S#{idx:02d}" for idx in range(1, count + 1)]


def _final_sample_numbers(data: ExtractedFinalReportData) -> list[str]:
    if data.sample_numbers:
        return data.sample_numbers
    count = data.tested_sample_count or data.requested_sample_count or len(data.sample_results or [])
    if count:
        return [_sample_no_for_index(data.seal_part_no, idx) for idx in range(1, count + 1)]
    return [result.sample_no for result in data.sample_results or []] or ["S#01"]


def _results_for_samples(data: ExtractedFinalReportData, samples: list[str]) -> list[SampleResult]:
    by_suffix = {_sample_suffix(result.sample_no): result for result in data.sample_results or []}
    results: list[SampleResult] = []
    for sample in samples:
        suffix = _sample_suffix(sample)
        results.append(by_suffix.get(suffix, SampleResult(sample, data.target_hours, data.actual_hours, "", "")))
    return results


def _sample_column_label(sample: str, part_no: str) -> str:
    suffix = _sample_suffix(sample)
    sample_part = _part_no_from_sample(sample) or part_no
    return f"{sample_part}-{suffix}" if sample_part else suffix


def _sample_no_for_index(part_no: str, idx: int) -> str:
    suffix = f"S#{idx:02d}"
    return f"{part_no} - {suffix}" if part_no else suffix


def _part_no_from_sample(sample: str) -> str:
    match = re.search(r"\b\d{5,8}\b", sample or "")
    return match.group(0) if match else ""


def _hours_text(value: str) -> str:
    number = _first_number(value)
    return f"{number} Hours" if number is not None else value


def _hours_text_short(value: str) -> str:
    number = _first_number(value)
    return f"{number} Hrs" if number is not None else value


def _text_after_heading(text: str, heading: str, max_lines: int = 20) -> str:
    # Match the heading loosely so variants like "Section - C", "Section: C"
    # or "SECTION–C" are still recognised, not just an exact "Section C".
    flexible = re.escape(heading).replace(r"\ ", r"\s*[-–—:]*\s*")
    pattern = re.compile(rf"^\s*{flexible}\b", re.IGNORECASE)
    lines = (text or "").splitlines()
    for idx, line in enumerate(lines):
        if pattern.search(line.strip()):
            return "\n".join(lines[idx + 1:idx + 1 + max_lines])
    return ""


def _slice_between(text: str, start_heading: str, end_headings: list[str]) -> str:
    """Return the lines strictly between a heading line (e.g. "Shaft:") and
    whichever of end_headings comes first, so labels shared by multiple
    sections (Diameter, Roughness, ...) can be looked up unambiguously."""
    lines = (text or "").splitlines()
    start_idx = None
    for idx, line in enumerate(lines):
        if line.strip().rstrip(":").casefold() == start_heading.casefold():
            start_idx = idx + 1
            break
    if start_idx is None:
        return ""
    end_set = {heading.casefold() for heading in end_headings}
    end_idx = len(lines)
    for idx in range(start_idx, len(lines)):
        if lines[idx].strip().rstrip(":").casefold() in end_set:
            end_idx = idx
            break
    return "\n".join(lines[start_idx:end_idx])


def _extract_dashed_value(text: str, labels: list[str]) -> str:
    """Like _extract_label_value, but the project specification also uses
    en/em dashes as a "label – value" separator (e.g. "Fluid Level –
    Shaft Centerline"), not just a colon."""
    for label in labels:
        pattern = re.compile(rf"\b{re.escape(label)}\s*[:\-–—]\s*([^\n]+)", re.IGNORECASE)
        match = pattern.search(text or "")
        if match:
            value = match.group(1).strip()
            if value:
                return value
    return ""


_SECTION_C_ROW_LABELS = ["DRO", "STBM", "Seal Cock", "Reciprocation"]
_ROMAN_NUMERAL_RE = r"(I{1,3}|IV|V)"


def _extract_section_c_rows(text: str) -> dict[str, dict[str, str]]:
    """Read one row per sample out of the Inspection & Execution sheet's
    Section C (Test Log): Start Date, End Date, Position (the
    "II - Left" style GB1GTMC011 slot code), and the handwritten
    DRO/STBM/Seal Cock/Reciprocation figures. Rows are labelled "#1".."#6"
    in the sheet. Best-effort: these are handwritten values that go
    through OCR, so column association is inferred from the fixed
    left-to-right order of the printed columns rather than any
    positional/layout data."""
    section = (
        _slice_between(text, "C. TEST LOG", ["D. SPECIFIC REQUESTOR REQUIREMENTS", "E. OBSERVATIONS / ADDITIONAL INFORMATION DURING EXECUTION", "F. TEST CLOSURE"])
        or _text_after_heading(text, "Section C", max_lines=60)
        or _text_after_heading(text, "C. TEST LOG", max_lines=60)
    )
    rows: dict[str, dict[str, str]] = {}
    for line in section.splitlines():
        match = re.match(r"^\s*[A-Za-z]*#\s*0?(\d{1,2})\b(.*)$", line.strip())
        if not match:
            continue
        sample_key = f"S#{int(match.group(1)):02d}"
        rest = match.group(2)

        row: dict[str, str] = {}
        dates = re.findall(r"\b\d{1,2}[/.\-]\w{2,9}[/.\-]\d{2,4}\b", rest)
        if len(dates) >= 1:
            row["Start Date"] = dates[0]
        if len(dates) >= 2:
            row["End Date"] = dates[1]

        position_match = re.search(rf"\b{_ROMAN_NUMERAL_RE}\s*[-–—]\s*(Left|Right)\b", rest, re.IGNORECASE)
        if position_match:
            row["Position"] = f"{position_match.group(1).upper()} - {position_match.group(2).title()}"

        decimals = re.findall(r"\b0?\.\d{2,3}\b", rest)
        for label, value in zip(_SECTION_C_ROW_LABELS, decimals):
            row[label] = value

        if row:
            rows[sample_key] = row
    return rows


_ROMAN_TO_MACHINE_NUMBER = {"I": "01", "II": "02", "III": "03", "IV": "04", "V": "05"}


def _machine_identification(position_code: str) -> tuple[str, str]:
    """Turn a Section C "Position (Only GB1GTMC011)" entry like "II - Left"
    into (machine_identification, position): "II" is machine #02, reported
    as "Hot Seal Endurance machine #02 (GB1GTMC011-II)"; the printed
    Left/Right is reported as the sample's Position as-is."""
    match = re.match(rf"\s*{_ROMAN_NUMERAL_RE}\s*[-–—]\s*(Left|Right)", position_code or "", re.IGNORECASE)
    if not match:
        return "", ""
    roman = match.group(1).upper()
    position = match.group(2).title()
    number = _ROMAN_TO_MACHINE_NUMBER.get(roman)
    if not number:
        return "", position
    return f"Hot Seal Endurance machine #{number} (GB1GTMC011-{roman})", position


def _build_test_dates(data: ExtractedFinalReportData, inspection_text: str) -> list[TestDateEntry]:
    section_c = _extract_section_c_rows(inspection_text)
    entries: list[TestDateEntry] = []
    for sample in _final_sample_numbers(data):
        row = section_c.get(_sample_suffix(sample), {})
        machine_id, position = _machine_identification(row.get("Position", ""))
        entries.append(
            TestDateEntry(
                sample_no=sample,
                start_date=_format_long_date(row.get("Start Date", "")),
                end_date=_format_long_date(row.get("End Date", "")),
                position=position,
                machine_identification=machine_id,
            )
        )
    return entries


_ORDINAL_SUFFIXES = {1: "st", 2: "nd", 3: "rd"}


def _ordinal(day: int) -> str:
    suffix = "th" if 10 <= day % 100 <= 20 else _ORDINAL_SUFFIXES.get(day % 10, "th")
    return f"{day}{suffix}"


def _format_long_date(value: str) -> str:
    """Turn a short handwritten date like "15-Jan-25" or "15/01/2025" into
    the "15th January 2025" style used in the Final Report's Date the
    test was performed table."""
    value = (value or "").strip()
    if not value:
        return ""
    for fmt in ("%d-%b-%y", "%d-%b-%Y", "%d/%m/%y", "%d/%m/%Y", "%d.%m.%y", "%d.%m.%Y", "%d-%m-%y", "%d-%m-%Y"):
        try:
            parsed = dt.datetime.strptime(value, fmt)
        except ValueError:
            continue
        return f"{_ordinal(parsed.day)} {parsed.strftime('%B %Y')}"
    return value


_TOOLING_ROW_LABELS = ["OD (mm)", "Width (mm)", "Roughness, Ra (µm)", "Hardness (HRC)", "Lead", "ID (mm)"]


def _extract_tooling_measurements(text: str) -> dict[str, dict[str, str]]:
    """Read the "A. Tooling Measurement (Shaft & Bore Plate)" grid: each
    row (OD, Roughness, Hardness, ...) has one handwritten value per
    tooling column (#1..#6). Returns {row_label: {tooling_column: value}}.
    "Roughness, Ra (µm)" appears once for the shaft and once for the bore
    plate, so the second occurrence is keyed with a " [Bore]" suffix."""
    section = _slice_between(
        text,
        "A. TOOLING MEASUREMENT (SHAFT & BORE PLATE)",
        ["B. SEAL MEASUREMENT", "B. SEAL MEASUREMENT — PRE-TEST & POST-TEST", "Inspected By / Date"],
    ) or text

    result: dict[str, dict[str, str]] = {}
    seen_roughness = False
    for line in section.splitlines():
        stripped = line.strip()
        for label in _TOOLING_ROW_LABELS:
            if not stripped.casefold().startswith(label.casefold()):
                continue
            rest = stripped[len(label):]
            values = re.findall(r"\d+\.\d+|\d+", rest)
            if not values:
                break
            key = label
            if label == "Roughness, Ra (µm)":
                key = "Roughness, Ra (µm) [Bore]" if seen_roughness else "Roughness, Ra (µm)"
                seen_roughness = True
            result[key] = {str(idx): value for idx, value in enumerate(values, start=1)}
            break
    return result


def _extract_seal_used_mapping(text: str, side: str) -> dict[str, list[str]]:
    """Parse the "Seal used (Mention seal No. respective to shaft/Bore)"
    row: the technician writes which sample number(s) were tested with
    each tooling column, e.g. "1, 2" under column #1 and "3" under column
    #2. Returns {tooling_column: [sample_number, ...]}."""
    for line in (text or "").splitlines():
        stripped = line.strip()
        lowered = stripped.casefold()
        if "seal used" not in lowered or f"respective to {side}".casefold() not in lowered:
            continue
        rest = re.sub(r"(?i)seal used\s*\(mention seal no\.?\s*respective to \w+\)?", "", stripped).strip()
        groups = re.findall(r"\d+(?:\s*,\s*\d+)*", rest)
        mapping: dict[str, list[str]] = {}
        for column_idx, group in enumerate(groups, start=1):
            sample_numbers = [token.strip() for token in group.split(",") if token.strip()]
            if sample_numbers:
                mapping[str(column_idx)] = sample_numbers
        return mapping
    return {}


def _values_by_sample(column_to_samples: dict[str, list[str]], column_to_value: dict[str, str]) -> dict[str, str]:
    """Combine a tooling-column -> sample-numbers mapping with a
    tooling-column -> measured-value mapping into sample -> value."""
    result: dict[str, str] = {}
    for column, sample_numbers in column_to_samples.items():
        value = column_to_value.get(column)
        if not value:
            continue
        for sample_number in sample_numbers:
            try:
                result[f"S#{int(sample_number):02d}"] = value
            except ValueError:
                continue
    return result


def _combined_seal_used_map(text: str) -> dict[str, list[str]]:
    """Section B's per-column sample numbering is the same "Seal Used
    (Mention seal No. respective to shaft/bore)" pair of rows used for
    Section A's tooling grid - the tooling column is a physical rig
    position, so it names the same sample regardless of which row is
    read. Shaft takes precedence where both rows name a column, purely to
    have a single deterministic source; in practice they should agree."""
    shaft_map = _extract_seal_used_mapping(text, "shaft")
    bore_map = _extract_seal_used_mapping(text, "bore")
    combined = dict(bore_map)
    combined.update(shaft_map)
    return combined


_SECTION_B_END_HEADINGS = ["C. TEST LOG", "Section C"]


def _extract_section_b_rows(text: str, column_count: int) -> tuple[list[tuple[str, str, list[str]]], str]:
    """Read the handwritten "B. Seal Measurement - Pre-Test & Post-Test"
    grid. Each filled-in row is a measurement description, an optional
    hand-written Spec value, then one pre-test value and one post-test
    value per tooling column, pre first then post (e.g. "54.050 54.143
    54.024 54.079" for two columns). Which description labels appear -
    and whether a Spec was written at all - differs project to project,
    so no fixed label list or fixed Spec presence is assumed (unlike
    Section A's tooling grid).

    The row only tells us the *count* of numbers, not which ones are the
    Spec versus the readings, since the Spec may be a single nominal
    value or a "nominal +/- tolerance" pair written free-hand. `column_count`
    (the number of tooling columns the "Seal used" row named) tells us how
    many numbers the pre/post readings must occupy - 2 per column - so any
    numbers beyond that, from the left, are the Spec, captured as the
    original substring (not just the extracted digits) so a hand-written
    "82.32 +/- 0.07" or "54.00-54.05" survives as written.

    The "Instrument Used/ID" row and the "Seal used" mapping rows are
    deliberately skipped - the former identifies the physical gauge, not
    a seal measurement, and has no place in the final report; the latter
    is read separately by _combined_seal_used_map. Any "Remarks" line is
    captured verbatim and returned separately as free text."""
    section = (
        _slice_between(text, "B. SEAL MEASUREMENT — PRE-TEST & POST-TEST", _SECTION_B_END_HEADINGS)
        or _slice_between(text, "B. SEAL MEASUREMENT", _SECTION_B_END_HEADINGS)
        or _text_after_heading(text, "B. Seal Measurement", max_lines=80)
    )
    rows: list[tuple[str, str, list[str]]] = []
    remarks: list[str] = []
    for line in section.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        lowered = stripped.casefold()
        if "instrument" in lowered and ("used" in lowered or "id" in lowered):
            continue
        if "seal used" in lowered:
            continue
        if "remark" in lowered:
            note = re.sub(r"(?i)^.*?remarks?\s*[:\-]?\s*", "", stripped).strip()
            if note:
                remarks.append(note)
            continue
        match = re.match(r"^(?P<label>[A-Za-z][^\d]*?)\s*[:\-]?\s*(?P<rest>\d.*)$", stripped)
        if not match:
            continue
        label = match.group("label").strip().rstrip(":-").strip()
        rest = match.group("rest")
        number_matches = list(re.finditer(r"-?\d+\.\d+|-?\d+", rest))
        if not label or len(number_matches) < 2:
            continue

        measurement_count = 2 * column_count if column_count > 0 else len(number_matches) - (len(number_matches) % 2)
        spec_token_count = max(len(number_matches) - measurement_count, 0)
        spec_text = ""
        if spec_token_count:
            spec_text = rest[:number_matches[spec_token_count - 1].end()].strip()
        numbers = [m.group(0) for m in number_matches[spec_token_count:]]
        rows.append((label, spec_text, numbers))
    return rows, " ".join(remarks)


def _pre_post_pairs(numbers: list[str]) -> list[tuple[str, str]]:
    return [(numbers[i], numbers[i + 1]) for i in range(0, len(numbers) - 1, 2)]


def _extract_section_b_measurements(
    text: str, seal_used_map: dict[str, list[str]]
) -> tuple[dict[str, dict[str, tuple[str, str]]], dict[str, str], str]:
    """{description label: {sample suffix: (pre value, post value)}}, plus
    {description label: hand-written Spec text} and any Section B remarks.
    Built by pairing each row's pre/post-per-column values with the sample
    number the "Seal used" row assigned to that column."""
    column_count = max((int(column) for column in seal_used_map), default=0)
    rows, remarks_text = _extract_section_b_rows(text, column_count)
    measurements: dict[str, dict[str, tuple[str, str]]] = {}
    specs: dict[str, str] = {}
    for label, spec_text, numbers in rows:
        by_sample: dict[str, tuple[str, str]] = {}
        for column_idx, (pre_value, post_value) in enumerate(_pre_post_pairs(numbers), start=1):
            for sample_number in seal_used_map.get(str(column_idx), []):
                try:
                    suffix = f"S#{int(sample_number):02d}"
                except ValueError:
                    continue
                by_sample[suffix] = (pre_value, post_value)
        if by_sample:
            measurements[label] = by_sample
            if spec_text:
                specs[label] = spec_text
    return measurements, specs, remarks_text


def _project_no_from_filename(path: Path) -> str:
    return _first_match(PROJECT_NO_RE, path.name)


def _display_date(value: str) -> str:
    text = (value or "").strip()
    for fmt in ("%d/%m/%Y", "%Y-%m-%d"):
        try:
            return dt.datetime.strptime(text, fmt).strftime("%d %B %Y")
        except ValueError:
            pass
    return text or dt.date.today().strftime("%d %B %Y")


def _safe_filename(value: str) -> str:
    return re.sub(r"[^A-Za-z0-9 ._()-]+", "_", value).strip(" ._") or "Final report"


def _paragraph_index(doc: Document, text: str, start: int = 0) -> int | None:
    for idx in range(start, len(doc.paragraphs)):
        if doc.paragraphs[idx].text.strip().casefold() == text.casefold():
            return idx
    return None
